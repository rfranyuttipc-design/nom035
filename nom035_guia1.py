"""
NOM-035-STPS-2018 | RFRANYUTTI, CONCIENCIA VERDE Y LABORAL S.C.
================================================================
MÓDULO 0 + 1  —  Panel Operativo + Cuestionario Virtual
GUÍA I        —  Identificación de Trabajadores Sujetos a
                 Acontecimientos Traumáticos Severos

CORRECCIONES v2:
- RFRANYUTTI eliminado de clientes (es el prestador)
- Rangos de edad 15-19 … 70+
- Estado civil añadido
- Nivel de estudios: Sin formación → Doctorado
- Tipo de jornada reemplaza turno
- Tiempo en puesto y experiencia: 8 rangos cada uno
- BORRAR DATOS corregido (limpia campos realmente)
- Texto de confirmación exacto de la norma
- Preguntas oficiales NOM-035 Guía I (S1:6, S2:2, S3:7, S4:5)
- Lógica de resultado corregida a criterios exactos NOM-035
- PermissionError resuelto con reintento automático
"""

import streamlit as st
import pandas as pd
import os, re, time, sys, sqlite3, threading
from datetime import datetime

# ── Módulo de reportes (embebido) ──────────────────────────────────────────
import io, textwrap

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from openpyxl import load_workbook
from openpyxl.styles import (
    PatternFill, Font, Alignment, Border, Side, GradientFill
)
from openpyxl.chart import BarChart, PieChart, Reference
from openpyxl.chart.series import DataPoint
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paleta de colores RFRANYUTTI ────────────────────────────────────────────
VERDE        = "4B694E"
VERDE_CLARO  = "D6E4D8"
VERDE_MID    = "6A9370"
ROJO         = "A20000"
ROJO_CLARO   = "FDDEDE"
GRIS_HEADER  = "F4F4F4"
BLANCO       = "FFFFFF"
AMARILLO     = "FFF8DC"
AZUL_CLARO   = "EBF3FB"

# ─── Colores matplotlib ───────────────────────────────────────────────────────
C_VERDE  = "#4B694E"
C_VERDE2 = "#6A9370"
C_ROJO   = "#A20000"
C_AMBAR  = "#C8A600"
C_GRIS   = "#B4B4B4"
C_AZUL   = "#69A2D8"

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def _side(color="D0D0D0", style="thin"):
    return Side(border_style=style, color=color)

def _border(color="D0D0D0"):
    s = _side(color)
    return Border(left=s, right=s, top=s, bottom=s)

def _fill(hex_color):
    return PatternFill("solid", fgColor=hex_color)

def _font(bold=False, size=10, color="1A1A1A", name="Calibri"):
    return Font(bold=bold, size=size, color=color, name=name)

def _align(h="left", v="center", wrap=False):
    return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

def _auto_width(ws, min_w=8, max_w=40):
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            try:
                val = str(cell.value) if cell.value else ""
                max_len = max(max_len, len(val))
            except:
                pass
        ws.column_dimensions[col_letter].width = min(max(max_len + 2, min_w), max_w)

def _grafica_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=140, bbox_inches="tight",
                facecolor="white")
    buf.seek(0)
    plt.close(fig)
    return buf

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICAS MATPLOTLIB (se usan en Word y como imagen en Excel)
# ══════════════════════════════════════════════════════════════════════════════
def _graf_genero(df) -> io.BytesIO:
    conteo = df["Sexo"].value_counts()
    fig, ax = plt.subplots(figsize=(4.5, 3.5))
    colores = [C_VERDE, C_VERDE2, C_GRIS]
    wedges, texts, autotexts = ax.pie(
        conteo.values, labels=conteo.index,
        autopct="%1.1f%%", colors=colores[:len(conteo)],
        startangle=90, wedgeprops=dict(edgecolor="white", linewidth=1.5)
    )
    for t in texts:     t.set_fontsize(9)
    for a in autotexts: a.set_fontsize(8); a.set_color("white"); a.set_fontweight("bold")
    ax.set_title("Participación por Género", fontsize=11, fontweight="bold", pad=12)
    fig.tight_layout()
    return _grafica_bytes(fig)

def _graf_area(df) -> io.BytesIO:
    conteo = df["Área"].value_counts().head(10)
    fig, ax = plt.subplots(figsize=(6, max(3, len(conteo) * 0.55)))
    bars = ax.barh(conteo.index[::-1], conteo.values[::-1],
                   color=C_VERDE, edgecolor="white", height=0.6)
    for bar, val in zip(bars, conteo.values[::-1]):
        ax.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height()/2,
                str(val), va="center", fontsize=8, fontweight="bold", color=C_VERDE)
    ax.set_xlabel("Número de trabajadores", fontsize=9)
    ax.set_title("Participación por Área", fontsize=11, fontweight="bold", pad=10)
    ax.spines[["top","right"]].set_visible(False)
    ax.tick_params(axis="y", labelsize=8)
    fig.tight_layout()
    return _grafica_bytes(fig)

def _graf_riesgo(df) -> io.BytesIO:
    if "ATS" not in df.columns:
        return None
    req   = (df.get("Atención Clínica","No") == "Sí").sum()
    no_req = len(df) - req
    cats   = ["No requiere\nvaloraci\u00f3n", "Requiere\natenci\u00f3n cl\u00ednica"]
    vals   = [no_req, req]
    cols   = [C_VERDE, C_ROJO]
    fig, ax = plt.subplots(figsize=(5, 3.5))
    bars = ax.bar(cats, vals, color=cols, edgecolor="white", width=0.45)
    for bar, val in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1,
                str(val), ha="center", fontsize=10, fontweight="bold")
    ax.set_ylabel("Número de trabajadores", fontsize=9)
    ax.set_title("Resultado de Identificaci\u00f3n ATS", fontsize=11, fontweight="bold", pad=10)
    ax.spines[["top","right"]].set_visible(False)
    ax.set_ylim(0, max(vals) * 1.25 + 1)
    fig.tight_layout()
    return _grafica_bytes(fig)

def _graf_puesto(df) -> io.BytesIO:
    conteo = df["Puesto"].value_counts()
    fig, ax = plt.subplots(figsize=(5, 3.5))
    cols = [C_VERDE, C_VERDE2, C_AMBAR, C_AZUL, C_GRIS]
    ax.bar(conteo.index, conteo.values,
           color=cols[:len(conteo)], edgecolor="white", width=0.5)
    for i, val in enumerate(conteo.values):
        ax.text(i, val + 0.1, str(val), ha="center", fontsize=9, fontweight="bold")
    ax.set_ylabel("Trabajadores", fontsize=9)
    ax.set_title("Distribuci\u00f3n por Puesto", fontsize=11, fontweight="bold", pad=10)
    ax.spines[["top","right"]].set_visible(False)
    ax.tick_params(axis="x", labelsize=8)
    fig.tight_layout()
    return _grafica_bytes(fig)


# ══════════════════════════════════════════════════════════════════════════════
# EXCEL MEJORADO
# ══════════════════════════════════════════════════════════════════════════════
def generar_excel_mejorado(excel_path: str, cliente: str, razon: str) -> str:
    """Lee el Excel existente y produce uno nuevo con formato profesional."""
    if not os.path.exists(excel_path):
        return None

    df = pd.read_excel(excel_path)
    if df.empty:
        return None

    fecha_str = datetime.now().strftime("%Y%m%d_%H%M")
    out_path  = excel_path.replace(".xlsx", f"_reporte_{fecha_str}.xlsx")

    # Guardar datos en nueva hoja limpia
    with pd.ExcelWriter(out_path, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="Datos", index=False)

    wb = load_workbook(out_path)
    ws = wb["Datos"]

    # ── Estilos de encabezado ──────────────────────────────────────────────
    header_fill  = _fill(VERDE)
    header_font  = _font(bold=True, size=10, color=BLANCO)
    header_align = _align("center", "center")

    for cell in ws[1]:
        cell.fill  = header_fill
        cell.font  = header_font
        cell.alignment = header_align
        cell.border = _border(VERDE_MID)

    # ── Filas alternadas + formato condicional simple ──────────────────────
    COLS_SI_NO = {"ATS": None, "Atención Clínica": None}
    # Detectar índices de columnas especiales
    for idx, cell in enumerate(ws[1], 1):
        if cell.value in COLS_SI_NO:
            COLS_SI_NO[cell.value] = idx

    for row_idx, row in enumerate(ws.iter_rows(min_row=2), start=2):
        bg = VERDE_CLARO if row_idx % 2 == 0 else BLANCO
        for cell in row:
            cell.fill      = _fill(bg)
            cell.font      = _font(size=9)
            cell.alignment = _align("left", "center", wrap=True)
            cell.border    = _border()

        # Color especial para columnas Sí/No
        for col_name, col_idx in COLS_SI_NO.items():
            if col_idx:
                c = ws.cell(row=row_idx, column=col_idx)
                if c.value == "Sí":
                    c.fill = _fill(ROJO_CLARO)
                    c.font = _font(bold=True, size=9, color=ROJO)
                elif c.value == "No":
                    c.fill = _fill(VERDE_CLARO)
                    c.font = _font(bold=True, size=9, color=VERDE)

    # ── Ancho de columnas + congelar encabezado ────────────────────────────
    _auto_width(ws)
    ws.freeze_panes = "A2"

    # ── Hoja RESUMEN ──────────────────────────────────────────────────────
    ws_res = wb.create_sheet("Resumen")
    total     = len(df)
    fem       = (df["Sexo"] == "Femenino").sum() if "Sexo" in df.columns else 0
    masc      = (df["Sexo"] == "Masculino").sum() if "Sexo" in df.columns else 0
    ats_si    = (df["ATS"] == "Sí").sum() if "ATS" in df.columns else 0
    atencion  = (df.get("Atención Clínica","") == "Sí").sum()

    resumen_data = [
        ("INDICADOR",                     "VALOR"),
        ("Total de trabajadores evaluados", total),
        ("Femenino",                        fem),
        ("Masculino",                       masc),
        ("ATS identificado",                ats_si),
        ("No requiere valoración clínica",  total - ats_si),
        ("Requiere atención clínica",       atencion),
    ]

    ws_res["A1"] = f"RESUMEN EJECUTIVO — {razon}"
    ws_res["A1"].font  = _font(bold=True, size=13, color=VERDE)
    ws_res["A1"].alignment = _align("left", "center")
    ws_res.merge_cells("A1:B1")
    ws_res["A2"] = f"Generado: {datetime.now().strftime('%d %b %Y %H:%M')}"
    ws_res["A2"].font  = _font(size=9, color="888888")
    ws_res.merge_cells("A2:B2")

    for i, (label, val) in enumerate(resumen_data, start=4):
        c_label = ws_res.cell(row=i, column=1, value=label)
        c_val   = ws_res.cell(row=i, column=2, value=val)
        if i == 4:  # encabezado
            for c in [c_label, c_val]:
                c.fill  = _fill(VERDE)
                c.font  = _font(bold=True, size=10, color=BLANCO)
                c.alignment = _align("center", "center")
        else:
            bg = GRIS_HEADER if i % 2 == 0 else BLANCO
            for c in [c_label, c_val]:
                c.fill      = _fill(bg)
                c.font      = _font(size=10)
                c.alignment = _align("left", "center")
                c.border    = _border()

    ws_res.column_dimensions["A"].width = 38
    ws_res.column_dimensions["B"].width = 14

    # ── Insertar gráficas como imágenes en hoja Gráficas ──────────────────
    ws_graf = wb.create_sheet("Gráficas")
    ws_graf["A1"] = "GRÁFICAS DE RESULTADOS — NOM-035-STPS-2018"
    ws_graf["A1"].font = _font(bold=True, size=13, color=VERDE)
    ws_graf.merge_cells("A1:N1")

    graficas = [
        ("Género",  _graf_genero(df),  "A3"),
        ("Área",    _graf_area(df),    "I3"),
        ("Riesgo",  _graf_riesgo(df),  "A28"),
        ("Puesto",  _graf_puesto(df),  "I28"),
    ]
    from openpyxl.drawing.image import Image as XLImage
    for titulo, buf, celda in graficas:
        if buf:
            buf.seek(0)
            img = XLImage(buf)
            img.width  = 340
            img.height = 240
            ws_graf.add_image(img, celda)

    # ── Hoja de Analíticas avanzadas ─────────────────────────────────────
    analiticas = _calcular_analiticas(df)
    if analiticas:
        _hoja_analiticas(wb, df, analiticas)

    wb.save(out_path)
    return out_path


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO WORD
# ══════════════════════════════════════════════════════════════════════════════
def _set_cell_bg(cell, hex_color):
    """Sets table cell background color via XML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def _set_cell_border(cell, color="D0D0D0"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top","left","bottom","right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def _par(doc, text="", bold=False, size=10, color=None,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold           = bold
        run.font.name      = "Calibri"
        run.font.size      = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
    return p

_fig_num_word = [0]  # contador global de figuras en Word

def _insertar_imagen_word(doc, buf, width_cm=14, caption=None):
    if not buf:
        return
    buf.seek(0)
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(4)
    p_img.paragraph_format.space_after  = Pt(2)
    p_img.add_run().add_picture(buf, width=Cm(width_cm))
    if caption:
        _fig_num_word[0] += 1
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        r_cp = cp.add_run(f"Figura {_fig_num_word[0]}. {caption}")
        r_cp.font.name      = "Calibri"
        r_cp.font.size      = Pt(8.5)
        r_cp.italic         = True
        r_cp.font.color.rgb = RGBColor(0x60,0x60,0x60)

def generar_word(excel_path: str, cliente: str, razon: str,
                 logo_rf: str = None, logo_cliente: str = None) -> str:
    """Genera informe Word NOM-035 Guía I.
    - Calibri 11 en todo el documento
    - Texto justificado
    - Encabezado estilo NOM-004 desde pág 2
    - Figuras enumeradas automáticamente
    - Pie: DOCUMENTO ORIGINAL Calibri 8 RGB(96,96,96)
    """
    if not os.path.exists(excel_path):
        return None
    df = pd.read_excel(excel_path)
    if df.empty:
        return None

    fecha_str = datetime.now().strftime("%Y%m%d_%H%M")
    _meses = {"January":"enero","February":"febrero","March":"marzo",
              "April":"abril","May":"mayo","June":"junio","July":"julio",
              "August":"agosto","September":"septiembre","October":"octubre",
              "November":"noviembre","December":"diciembre"}
    _fd       = datetime.now()
    fecha_es  = f"{_fd.day} de {_meses.get(_fd.strftime('%B'),_fd.strftime('%B'))} de {_fd.year}"
    out_path  = excel_path.replace(".xlsx", f"_informe_{fecha_str}.docx")

    doc     = Document()
    _fig_n  = [0]   # contador de figuras

    # ── Márgenes ─────────────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)
        sec.header_distance = Cm(1.0)
        sec.footer_distance = Cm(1.0)

    # ── Estilo global: Calibri 11 justificado ─────────────────────────────────
    nml = doc.styles["Normal"]
    nml.font.name  = "Calibri"
    nml.font.size  = Pt(11)
    nml.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── Helpers ───────────────────────────────────────────────────────────────
    def _run(paragraph, text, bold=False, size=11, color_hex=None, italic=False):
        r = paragraph.add_run(text)
        r.bold        = bold
        r.italic      = italic
        r.font.name   = "Calibri"
        r.font.size   = Pt(size)
        if color_hex:
            r.font.color.rgb = RGBColor.from_string(color_hex)
        return r

    def _p(text="", bold=False, size=11, color_hex=None,
           align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6, italic=False):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        if text:
            _run(p, text, bold=bold, size=size, color_hex=color_hex, italic=italic)
        return p

    def _section_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "8")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), VERDE)
        pBdr.append(bot)
        pPr.append(pBdr)
        _run(p, text, bold=True, size=13, color_hex=VERDE)
        return p

    def _subsection(text, color_hex=VERDE):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(3)
        _run(p, text, bold=True, size=11, color_hex=color_hex)
        return p

    def _figura(buf, caption, width_cm=13):
        """Inserta imagen + Figura N. Caption (Calibri 9 cursiva gris)"""
        if not buf: return
        buf.seek(0)
        _fig_n[0] += 1
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after  = Pt(2)
        p_img.add_run().add_picture(buf, width=Cm(width_cm))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        _run(p_cap, f"Figura {_fig_n[0]}. {caption}",
             size=9, italic=True, color_hex="606060")

    def _set_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _set_bdr(cell, color="4B694E"):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcB  = OxmlElement("w:tcBorders")
        for side in ["top","left","bottom","right"]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            tcB.append(el)
        tcPr.append(tcB)

    def _valign(cell, val="center"):
        tcPr  = cell._tc.get_or_add_tcPr()
        vAl   = OxmlElement("w:vAlign")
        vAl.set(qn("w:val"), val)
        tcPr.append(vAl)

    def _tbl_cell(cell, text, bold=False, size=10, color_hex=None,
                  align=WD_ALIGN_PARAGRAPH.CENTER, bg=None, bdr="4B694E",
                  sb=3, sa=3):
        if bg: _set_bg(cell, bg)
        _set_bdr(cell, bdr)
        _valign(cell)
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        _run(p, text, bold=bold, size=size, color_hex=color_hex)

    def _col_width(cell, cm):
        tcPr = cell._tc.get_or_add_tcPr()
        tcW  = OxmlElement("w:tcW")
        tcW.set(qn("w:w"),    str(int(cm * 567)))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

    # ══════════════════════════════════════════════════════════════════════
    # ENCABEZADO NOM-004 (desde pág 2)
    # Estructura:
    #   ┌────────┬──────────────────────────────┐
    #   │        │ Empresa: NOMBRE EN NEGRITAS  │
    #   │  LOGO  │ TÍTULO DOCUMENTO EN MAYÚS    │
    #   │        ├────────────────┬─────────────┤
    #   │        │ Fecha emisión  │  No. Página │
    #   └────────┴────────────────┴─────────────┘
    # ══════════════════════════════════════════════════════════════════════
    def _build_header(section):
        section.different_first_page_header_footer = True

        # Primera página: encabezado vacío (portada)
        fph = section.first_page_header
        if not fph.paragraphs:
            fph.add_paragraph()
        fph.paragraphs[0].text = ""

        # Pie de página (todas las páginas)
        for footer_obj in [section.footer, section.first_page_footer]:
            for p in footer_obj.paragraphs:
                for r in p.runs: r.text = ""
            fp = footer_obj.paragraphs[0] if footer_obj.paragraphs else footer_obj.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fp.paragraph_format.space_before = Pt(0)
            fp.paragraph_format.space_after  = Pt(0)
            pPr  = fp._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            top_el = OxmlElement("w:top")
            top_el.set(qn("w:val"),   "single")
            top_el.set(qn("w:sz"),    "4")
            top_el.set(qn("w:space"), "1")
            top_el.set(qn("w:color"), "D0D0D0")
            pBdr.append(top_el)
            pPr.append(pBdr)
            r_pie = fp.add_run("DOCUMENTO ORIGINAL")
            r_pie.font.name      = "Calibri"
            r_pie.font.size      = Pt(8)
            r_pie.font.color.rgb = RGBColor(96, 96, 96)

        # Encabezado desde pág 2
        hdr = section.header
        for p in hdr.paragraphs:
            for r in p.runs: r.text = ""

        # Tabla 3 filas × 3 cols (col 0 = logo, col 1-2 = datos)
        # Después fusionamos col 0 filas 0-2 y col 1-2 fila 2
        tbl = hdr.add_table(rows=3, cols=3, width=Cm(16))
        tbl.style = "Table Grid"

        # Anchos: logo=3cm, centro=10cm, pag=3cm
        COL_W = [3, 10, 3]
        for ri in range(3):
            for ci in range(3):
                _col_width(tbl.cell(ri, ci), COL_W[ci])

        # Fusionar col 0 (logo: span 3 filas)
        tbl.cell(0,0).merge(tbl.cell(2,0))
        # Fusionar col 2 filas 0-1 (solo fila 2 tiene nº pág separado)
        tbl.cell(0,2).merge(tbl.cell(1,2))

        # Col 0: Logo corporativo centrado
        c_logo = tbl.cell(0,0)
        _set_bg(c_logo, BLANCO)
        _set_bdr(c_logo, VERDE)
        _valign(c_logo, "center")
        p_l = c_logo.paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_l.paragraph_format.space_before = Pt(3)
        p_l.paragraph_format.space_after  = Pt(3)
        if logo_cliente and os.path.exists(logo_cliente):
            try:
                p_l.add_run().add_picture(logo_cliente, width=Cm(2.6))
            except:
                _run(p_l, cliente, bold=True, size=9, color_hex=VERDE)
        else:
            _run(p_l, cliente, bold=True, size=9, color_hex=VERDE)

        # Col 1 Fila 0: Empresa
        c_emp = tbl.cell(0,1)
        _set_bg(c_emp, BLANCO)
        _set_bdr(c_emp, VERDE)
        _valign(c_emp, "center")
        p_emp = c_emp.paragraphs[0]
        p_emp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_emp.paragraph_format.space_before = Pt(3)
        p_emp.paragraph_format.space_after  = Pt(3)
        _run(p_emp, "Empresa:  ", size=7.5, color_hex="888888")
        _run(p_emp, razon,        size=8,   bold=True, color_hex=VERDE)

        # Col 1 Fila 1: Título documento
        c_tit = tbl.cell(1,1)
        _set_bg(c_tit, VERDE_CLARO)
        _set_bdr(c_tit, VERDE)
        _valign(c_tit, "center")
        p_tit = c_tit.paragraphs[0]
        p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_tit.paragraph_format.space_before = Pt(3)
        p_tit.paragraph_format.space_after  = Pt(3)
        _run(p_tit,
             "NOM-035-STPS-2018 — GUÍA I  |  CUESTIONARIO PARA IDENTIFICAR "
             "A LOS TRABAJADORES SUJETOS A ACONTECIMIENTOS TRAUMÁTICOS SEVEROS",
             size=7.5, bold=True, color_hex=VERDE)

        # Col 1 Fila 2: Fecha de emisión
        c_fec = tbl.cell(2,1)
        _set_bg(c_fec, BLANCO)
        _set_bdr(c_fec, VERDE)
        _valign(c_fec, "center")
        p_fec = c_fec.paragraphs[0]
        p_fec.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_fec.paragraph_format.space_before = Pt(3)
        p_fec.paragraph_format.space_after  = Pt(3)
        _run(p_fec, "Fecha de emisión:  ", size=7.5, color_hex="888888")
        _run(p_fec, fecha_es,              size=7.5, bold=True)

        # Col 2 Fila 2: Nº Página
        c_pag = tbl.cell(2,2)
        _set_bg(c_pag, BLANCO)
        _set_bdr(c_pag, VERDE)
        _valign(c_pag, "center")
        p_pag = c_pag.paragraphs[0]
        p_pag.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_pag.paragraph_format.space_before = Pt(3)
        p_pag.paragraph_format.space_after  = Pt(3)
        _run(p_pag, "No. Página:  ", size=7.5, color_hex="888888")
        fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"),"begin")
        ins  = OxmlElement("w:instrText"); ins.text = " PAGE "
        fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"),"end")
        r_pg = p_pag.add_run()
        r_pg._r.append(fld1); r_pg._r.append(ins); r_pg._r.append(fld2)
        r_pg.font.name = "Calibri"; r_pg.font.size = Pt(8); r_pg.bold = True

    _build_header(doc.sections[0])

    # ══════════════════════════════════════════════════════════════════════
    # PORTADA
    # ══════════════════════════════════════════════════════════════════════
    # Logos en tabla centrada
    tbl_port = doc.add_table(rows=1, cols=3)
    tbl_port.alignment = WD_TABLE_ALIGNMENT.CENTER
    if logo_rf and os.path.exists(logo_rf):
        c = tbl_port.cell(0,0)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        c.paragraphs[0].add_run().add_picture(logo_rf, width=Cm(3.5))
    sep_r = tbl_port.cell(0,1).paragraphs[0].add_run("|")
    sep_r.font.color.rgb = RGBColor(0xCC,0xCC,0xCC)
    sep_r.font.size = Pt(28)
    tbl_port.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_cliente and os.path.exists(logo_cliente):
        c = tbl_port.cell(0,2)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        c.paragraphs[0].add_run().add_picture(logo_cliente, width=Cm(3.5))

    _p(sa=6)

    # Línea verde decorativa
    p_ln = doc.add_paragraph()
    p_ln.paragraph_format.space_after = Pt(0)
    pPr = p_ln._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"12")
    bot.set(qn("w:space"),"1");    bot.set(qn("w:color"),VERDE)
    pBdr.append(bot); pPr.append(pBdr)

    _p("NOM-035-STPS-2018", bold=True, size=22, color_hex=VERDE,
       align=WD_ALIGN_PARAGRAPH.CENTER, sb=14, sa=3)
    _p("GUÍA DE REFERENCIA I", bold=True, size=13, color_hex=VERDE_MID,
       align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=5)
    _p("Cuestionario para Identificar a los Trabajadores que Fueron Sujetos "
       "a Acontecimientos Traumáticos Severos",
       size=11, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=14)

    # Tabla datos portada
    tbl_dat = doc.add_table(rows=3, cols=2)
    tbl_dat.style = "Table Grid"
    tbl_dat.alignment = WD_TABLE_ALIGNMENT.CENTER
    datos_port = [
        ("Empresa:",   razon),
        ("Fecha:",     fecha_es),
        ("Elaboró:",   "RFRANYUTTI, CONCIENCIA VERDE Y LABORAL S.C."),
    ]
    for i,(lbl,val) in enumerate(datos_port):
        c0 = tbl_dat.cell(i,0); c1 = tbl_dat.cell(i,1)
        _set_bg(c0, VERDE_CLARO); _set_bdr(c0, VERDE)
        _set_bg(c1, BLANCO);      _set_bdr(c1, VERDE)
        _tbl_cell(c0, lbl, bold=True, size=11, color_hex=VERDE,
                  align=WD_ALIGN_PARAGRAPH.LEFT, bg=None)
        _tbl_cell(c1, val, size=11,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, bg=None)

    _p(sa=8)

    # Estadísticas portada
    total    = len(df)
    ats_si   = (df["ATS"]=="Sí").sum() if "ATS" in df.columns else 0
    atencion = (df.get("Atención Clínica","")=="Sí").sum()
    tbl_st   = doc.add_table(rows=1, cols=3)
    tbl_st.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(num,etiq,col) in enumerate([
        (str(total),    "Trabajadores evaluados",  VERDE),
        (str(ats_si),   "Con ATS identificado",    ROJO if ats_si>0 else VERDE),
        (str(atencion), "Requieren atención",       ROJO if atencion>0 else VERDE),
    ]):
        c = tbl_st.cell(0,i)
        _set_bg(c, col); _set_bdr(c, col)
        _valign(c, "center")
        pn = c.add_paragraph()
        pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(pn, num, bold=True, size=28, color_hex=BLANCO)
        pe = c.add_paragraph()
        pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pe.paragraph_format.space_after = Pt(4)
        _run(pe, etiq, size=9, color_hex=BLANCO)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECCIÓN 1 — GRÁFICAS
    # ══════════════════════════════════════════════════════════════════════
    _section_title("1. RESULTADOS GRÁFICOS")
    _p("Las siguientes gráficas presentan los resultados de la evaluación "
       "aplicada conforme a la Guía de Referencia I de la NOM-035-STPS-2018.",
       sa=8)

    # Género y ATS lado a lado
    tbl_g = doc.add_table(rows=1, cols=2)
    tbl_g.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci,(buf,cap) in enumerate([
        (_graf_genero(df), "Participación por Género"),
        (_graf_riesgo(df), "Resultado de Identificación ATS"),
    ]):
        c = tbl_g.cell(0,ci)
        _valign(c,"center")
        if buf:
            buf.seek(0)
            _fig_n[0] += 1
            pi = c.paragraphs[0]
            pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pi.add_run().add_picture(buf, width=Cm(7.5))
            pc = c.add_paragraph()
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pc.paragraph_format.space_after = Pt(6)
            _run(pc, f"Figura {_fig_n[0]}. {cap}", size=9, italic=True, color_hex="606060")
    _p(sa=4)

    _figura(_graf_area(df),   "Participación por Área",       width_cm=14)
    _figura(_graf_puesto(df), "Distribución por Puesto",      width_cm=14)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECCIÓN 2 — TABLA DE RESULTADOS
    # ══════════════════════════════════════════════════════════════════════
    _section_title("2. TABLA DE RESULTADOS POR TRABAJADOR")
    _p(f"Se presentan los resultados individuales de los {total} trabajadores "
       "evaluados, indicando folio, datos demográficos y resultado de la "
       "evaluación conforme a los criterios de la NOM-035-STPS-2018.", sa=8)

    COLS_W = ["Folio","Nombre","Sexo","Edad","Puesto","Área",
              "ATS","Atención Clínica","Nivel"]
    cols_d = [c for c in COLS_W if c in df.columns]
    tbl_r  = doc.add_table(rows=1, cols=len(cols_d))
    tbl_r.style = "Table Grid"
    tbl_r.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i,col in enumerate(cols_d):
        c = tbl_r.rows[0].cells[i]
        _tbl_cell(c, col, bold=True, size=10, color_hex=BLANCO, bg=VERDE)

    for ri,(_,row) in enumerate(df.iterrows()):
        cells = tbl_r.add_row().cells
        bg_r  = GRIS_HEADER if ri%2==0 else BLANCO
        for i,col in enumerate(cols_d):
            val = str(row.get(col,"")) if pd.notna(row.get(col,"")) else ""
            c   = cells[i]
            if col in ["ATS","Atención Clínica"] and val=="Sí":
                _tbl_cell(c, val, bold=True, size=10, color_hex=ROJO, bg=ROJO_CLARO)
            else:
                _tbl_cell(c, val, size=10, bg=bg_r)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECCIÓN 3 — RECOMENDACIONES
    # ══════════════════════════════════════════════════════════════════════
    _section_title("3. RECOMENDACIONES CONFORME A NOM-035-STPS-2018")
    _p("Con base en los resultados obtenidos, se emiten las siguientes "
       "recomendaciones de acuerdo con los criterios de la NOM-035-STPS-2018.",
       sa=8)

    recomendaciones = [
        ("Trabajadores NO identificados con ATS",               VERDE, [
            "No se requiere valoración clínica derivada de esta evaluación.",
            "Mantener ambiente laboral favorable y continuar con evaluaciones periódicas.",
            "Implementar medidas preventivas: comunicación abierta y programas de bienestar.",
        ]),
        ("Identificados CON ATS — Sin criterio de atención clínica", "C8A600", [
            "El trabajador vivió acontecimiento(s) traumático(s) pero sin reacciones clínicas significativas.",
            "Brindar seguimiento psicológico preventivo en plazo no mayor a 30 días.",
            "Informar a Recursos Humanos para el plan de acción NOM-035.",
        ]),
        ("Trabajadores que REQUIEREN atención clínica",         ROJO, [
            "Presenta sintomatología que requiere intervención profesional prioritaria.",
            "Canalizar DE INMEDIATO a evaluación psicológica o médica especializada.",
            "Documentar la canalización en el expediente del programa NOM-035.",
            "El empleador garantiza la atención sin represalias (Art. 7 NOM-035-STPS-2018).",
        ]),
        ("Obligaciones generales del empleador",                VERDE, [
            "Informar a los trabajadores sobre los factores de riesgo psicosocial.",
            "Llevar registros de resultados y acciones implementadas.",
            "Establecer medidas de control ante factores de riesgo identificados.",
            "Practicar exámenes médicos a trabajadores expuestos a ATS.",
        ]),
    ]

    for tit_r, col_r, puntos in recomendaciones:
        tbl_t = doc.add_table(rows=1, cols=1)
        tbl_t.alignment = WD_TABLE_ALIGNMENT.CENTER
        c_t = tbl_t.cell(0,0)
        _tbl_cell(c_t, tit_r, bold=True, size=11, color_hex=BLANCO,
                  bg=col_r, align=WD_ALIGN_PARAGRAPH.LEFT)
        for punto in puntos:
            pb = doc.add_paragraph(style="List Bullet")
            pb.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pb.paragraph_format.left_indent = Cm(0.5)
            pb.paragraph_format.space_after  = Pt(3)
            _run(pb, punto, size=11)
        _p(sa=4)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECCIÓN 4 — ANALÍTICAS AVANZADAS
    # ══════════════════════════════════════════════════════════════════════
    analiticas = _calcular_analiticas(df)
    if not analiticas:
        doc.save(out_path)
        return out_path

    a = analiticas
    _section_title("4. ANALÍTICAS AVANZADAS")
    _p("Esta sección presenta indicadores cuantitativos: índice de canalización, "
       "prevalencia de síntomas, segmentación por área y clasificación de urgencia.",
       sa=8)

    # 4.1 Índice de canalización
    _subsection("4.1 Índice de Canalización", ROJO)
    _p(f"Del total de {a['total']} trabajadores evaluados, {a['aten_total']} "
       f"({a['pct_canalizacion']}%) requieren ser canalizados a atención médica "
       "conforme a los criterios de la NOM-035-STPS-2018.", sa=6)

    # KPIs + gráfica lado a lado
    tbl_can = doc.add_table(rows=1, cols=2)
    tbl_can.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Gráfica izq
    c_gi = tbl_can.cell(0,0)
    buf_c = _graf_canalizacion(a)
    _fig_n[0] += 1
    if buf_c:
        buf_c.seek(0)
        pi = c_gi.paragraphs[0]
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi.add_run().add_picture(buf_c, width=Cm(7))
        pc = c_gi.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_after = Pt(4)
        _run(pc, f"Figura {_fig_n[0]}. Índice de Canalización",
             size=9, italic=True, color_hex="606060")
    # KPIs der
    c_kd = tbl_can.cell(0,1)
    _valign(c_kd, "center")
    kpis = [
        ("Total evaluados",           str(a["total"]),                 VERDE),
        ("Con ATS identificado",       str(a["ats_total"]),            ROJO if a["ats_total"]>0 else VERDE),
        ("Requieren atención clínica", str(a["aten_total"]),           ROJO if a["aten_total"]>0 else VERDE),
        ("Índice de canalización",     f"{a['pct_canalizacion']}%",    ROJO if a["pct_canalizacion"]>0 else VERDE),
    ]
    tbl_k = c_kd.add_table(rows=len(kpis), cols=2)
    tbl_k.style = "Table Grid"
    for ki,(lbl,val,col_k) in enumerate(kpis):
        ck0 = tbl_k.cell(ki,0); ck1 = tbl_k.cell(ki,1)
        _tbl_cell(ck0, lbl, size=11, bg=GRIS_HEADER, align=WD_ALIGN_PARAGRAPH.LEFT)
        _tbl_cell(ck1, val, bold=True, size=12, color_hex=col_k, bg=BLANCO)
    _p(sa=8)

    # 4.2 Prevalencia
    _subsection("4.2 Prevalencia de Síntomas", "69A2D8")
    _p(f"Entre los {a['ats_total']} trabajadores con ATS identificado, "
       f"el síntoma más prevalente es '{a['sintoma_dominante']}'.", sa=6)
    _figura(_graf_prevalencia(a), "Prevalencia de Síntomas por Categoría", 13)

    # 4.3 Segmentación
    if a["seg_area"]:
        _subsection("4.3 Segmentación por Área", "C8A600")
        _p(f"El área con mayor concentración de casos es '{a['foco_rojo_area']}'. "
           "Semáforo: 🔴 ≥50% — 🟡 25-49% — 🟢 <25%.", sa=6)
        _figura(_graf_segmentacion_area(a),
                "Segmentación por Área — Casos que Requieren Atención", 13)

    # 4.4 Urgencia
    _subsection("4.4 Urgencia de Intervención", ROJO)
    if a["criticos"]:
        _p(f"Se identificaron {a['n_criticos']} caso(s) CRÍTICO(S) y "
           f"{a['n_altos']} caso(s) de urgencia ALTA.", sa=6)
        _figura(_graf_urgencia(a), "Niveles de Urgencia de Intervención", 9)

        cols_urg = ["Folio","Nombre","Área","Puesto","Urgencia"]
        tbl_urg  = doc.add_table(rows=1, cols=len(cols_urg))
        tbl_urg.style = "Table Grid"
        tbl_urg.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i,h in enumerate(cols_urg):
            _tbl_cell(tbl_urg.rows[0].cells[i], h, bold=True, size=10,
                      color_hex=BLANCO, bg=VERDE)
        for caso in sorted(a["criticos"], key=lambda x: 0 if x["Urgencia"]=="CRÍTICO" else 1):
            rc = tbl_urg.add_row().cells
            bg_u = ROJO_CLARO if caso["Urgencia"]=="CRÍTICO" else "FFF8DC"
            for i,col in enumerate(cols_urg):
                _tbl_cell(rc[i], str(caso.get(col,"")), size=10,
                          bold=(caso["Urgencia"]=="CRÍTICO"),
                          color_hex=ROJO if caso["Urgencia"]=="CRÍTICO" else None,
                          bg=bg_u)
        _p(sa=6)
    else:
        _p("No se identificaron casos de urgencia crítica o alta.",
           bold=True, color_hex=VERDE, sa=4)

    # 4.5 Párrafo ejecutivo
    _subsection("4.5 Párrafo Ejecutivo de Cumplimiento Legal", VERDE)
    tbl_ej = doc.add_table(rows=1, cols=1)
    tbl_ej.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_ej = tbl_ej.cell(0,0)
    _set_bg(c_ej, VERDE_CLARO); _set_bdr(c_ej, VERDE)
    p_ej = c_ej.paragraphs[0]
    p_ej.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_ej.paragraph_format.space_before = Pt(6)
    p_ej.paragraph_format.space_after  = Pt(6)
    p_ej.paragraph_format.left_indent  = Cm(0.3)
    p_ej.paragraph_format.right_indent = Cm(0.3)
    _run(p_ej, a["parrafo_cumplimiento"], size=11)

    doc.save(out_path)
    return out_path

def _calcular_analiticas(df: pd.DataFrame) -> dict:
    """
    Calcula las 4 analíticas avanzadas a partir del DataFrame de resultados.
    Retorna un dict con todos los indicadores listos para usar en Excel y Word.
    """
    total = len(df)
    if total == 0:
        return {}

    # ── Columnas de respuestas por sección ────────────────────────────────
    cols_s1 = [c for c in df.columns if c.startswith("S1P")]
    cols_s2 = [c for c in df.columns if c.startswith("S2P")]
    cols_s3 = [c for c in df.columns if c.startswith("S3P")]
    cols_s4 = [c for c in df.columns if c.startswith("S4P")]

    ats_mask    = df["ATS"] == "Sí" if "ATS" in df.columns else pd.Series([False]*total)
    aten_mask   = df.get("Atención Clínica", pd.Series(["No"]*total)) == "Sí"

    ats_total   = int(ats_mask.sum())
    aten_total  = int(aten_mask.sum())
    no_aten     = total - aten_total

    # ── 1. ÍNDICE DE CANALIZACIÓN ─────────────────────────────────────────
    pct_canalizacion = round(aten_total / total * 100, 1) if total > 0 else 0
    pct_sin_atencion = round(100 - pct_canalizacion, 1)

    # ── 2. PREVALENCIA DE SÍNTOMAS (S2, S3, S4) ──────────────────────────
    # Solo entre trabajadores identificados con ATS
    df_ats = df[ats_mask] if ats_total > 0 else df

    def _positivos_seccion(cols):
        if not cols or df_ats.empty:
            return 0
        sub = df_ats[cols].apply(pd.to_numeric, errors="coerce").fillna(0)
        # Cada celda puede ser "Sí"/"No" o ya numérico
        # Reconvertir texto "Sí"→1 "No"→0
        for c in cols:
            if df_ats[c].dtype == object:
                sub[c] = df_ats[c].map({"Sí": 1, "No": 0}).fillna(0)
        # Trabajadores con al menos 1 "Sí" en la sección
        return int((sub.sum(axis=1) > 0).sum())

    prev_s2 = _positivos_seccion(cols_s2)   # Recuerdos persistentes
    prev_s3 = _positivos_seccion(cols_s3)   # Esfuerzo por evitar
    prev_s4 = _positivos_seccion(cols_s4)   # Afectación

    base_prev = max(ats_total, 1)
    pct_s2 = round(prev_s2 / base_prev * 100, 1)
    pct_s3 = round(prev_s3 / base_prev * 100, 1)
    pct_s4 = round(prev_s4 / base_prev * 100, 1)

    sintoma_dominante = max(
        [("Recuerdos persistentes (Secc. II)", pct_s2),
         ("Esfuerzo por evitar (Secc. III)",   pct_s3),
         ("Afectación (Secc. IV)",              pct_s4)],
        key=lambda x: x[1]
    )[0] if ats_total > 0 else "Sin ATS identificados"

    # ── 3. SEGMENTACIÓN POR ÁREA ──────────────────────────────────────────
    area_col = "Área" if "Área" in df.columns else None
    seg_area = {}
    foco_rojo_area = "N/D"
    if area_col and aten_total > 0:
        # Trabajadores con atención clínica por área
        seg_raw = df[aten_mask][area_col].value_counts()
        # Total por área
        total_area = df[area_col].value_counts()
        for area, n_aten in seg_raw.items():
            n_total = total_area.get(area, 1)
            seg_area[area] = {
                "total":    int(n_total),
                "atencion": int(n_aten),
                "pct":      round(n_aten / n_total * 100, 1),
            }
        # Área con mayor concentración de casos (mayor %)
        if seg_area:
            foco_rojo_area = max(seg_area.items(), key=lambda x: x[1]["pct"])[0]

    # Ordenar por % descendente
    seg_area_sorted = sorted(seg_area.items(), key=lambda x: x[1]["pct"], reverse=True)

    # ── 4. URGENCIA DE INTERVENCIÓN ──────────────────────────────────────
    # Casos críticos: ATS=Sí + al menos 1 Sí en S2 + ≥3 Sí en S3 + ≥2 Sí en S4
    def _cuenta_si(row, cols):
        return sum(1 for c in cols if str(row.get(c, "No")).strip() == "Sí")

    criticos = []
    for _, row in df.iterrows():
        if str(row.get("ATS", "No")) != "Sí":
            continue
        s2 = _cuenta_si(row, cols_s2)
        s3 = _cuenta_si(row, cols_s3)
        s4 = _cuenta_si(row, cols_s4)
        # Criterio de máxima urgencia: cumple los 3 criterios clínicos
        urgencia = "CRÍTICO" if (s2 >= 1 and s3 >= 3 and s4 >= 2) \
                   else "ALTO"   if (s2 >= 1 or s3 >= 3 or s4 >= 2) \
                   else "MEDIO"
        if urgencia in ("CRÍTICO", "ALTO"):
            criticos.append({
                "Folio":    row.get("Folio",""),
                "Nombre":   row.get("Nombre",""),
                "Área":     row.get("Área",""),
                "Puesto":   row.get("Puesto",""),
                "S2 (Sí)":  s2,
                "S3 (Sí)":  s3,
                "S4 (Sí)":  s4,
                "Urgencia": urgencia,
            })
    n_criticos = sum(1 for c in criticos if c["Urgencia"] == "CRÍTICO")
    n_altos    = sum(1 for c in criticos if c["Urgencia"] == "ALTO")

    # ── Párrafo ejecutivo de cumplimiento ────────────────────────────────
    parrafo_cumplimiento = (
        f"En cumplimiento de la NOM-035-STPS-2018, se aplicó el cuestionario de la "
        f"Guía de Referencia I a un total de {total} trabajador{'es' if total != 1 else ''} "
        f"de la organización. De ellos, {ats_total} ({round(ats_total/total*100,1) if total else 0}%) "
        f"reportaron haber estado expuestos a uno o más acontecimientos traumáticos severos "
        f"durante o con motivo del trabajo. Tras la aplicación de los criterios clínicos "
        f"establecidos en la norma, {aten_total} trabajador{'es requieren' if aten_total != 1 else ' requiere'} "
        f"ser remitido{'s' if aten_total != 1 else ''} a servicios médicos de forma inmediata, "
        f"lo que representa el {pct_canalizacion}% del total evaluado. "
        + (f"Se identificaron {n_criticos} caso{'s' if n_criticos != 1 else ''} de máxima "
           f"urgencia que cumplen simultáneamente todos los criterios clínicos de las "
           f"Secciones II, III y IV. " if n_criticos > 0 else "")
        + (f"El área con mayor concentración de casos que requieren atención es "
           f"'{foco_rojo_area}', con un {seg_area[foco_rojo_area]['pct']}% de sus "
           f"trabajadores evaluados en situación de riesgo. " if foco_rojo_area != "N/D" else "")
        + f"El síntoma más prevalente entre el personal afectado corresponde a "
          f"'{sintoma_dominante}'."
    )

    return {
        "total":                total,
        "ats_total":            ats_total,
        "aten_total":           aten_total,
        "no_aten":              no_aten,
        "pct_canalizacion":     pct_canalizacion,
        "pct_sin_atencion":     pct_sin_atencion,
        "prev_s2":              prev_s2,
        "prev_s3":              prev_s3,
        "prev_s4":              prev_s4,
        "pct_s2":               pct_s2,
        "pct_s3":               pct_s3,
        "pct_s4":               pct_s4,
        "sintoma_dominante":    sintoma_dominante,
        "seg_area":             seg_area_sorted,
        "foco_rojo_area":       foco_rojo_area,
        "criticos":             criticos,
        "n_criticos":           n_criticos,
        "n_altos":              n_altos,
        "parrafo_cumplimiento": parrafo_cumplimiento,
    }


# ─── Gráficas analíticas ─────────────────────────────────────────────────────

def _graf_canalizacion(a: dict) -> io.BytesIO:
    """Gráfica de dona — Índice de Canalización."""
    vals   = [a["aten_total"], a["no_aten"]]
    labels = [f"Requieren atención\n{a['pct_canalizacion']}%",
              f"Sin atención clínica\n{a['pct_sin_atencion']}%"]
    colors = [C_ROJO, C_VERDE]
    fig, ax = plt.subplots(figsize=(4.5, 3.8))
    wedges, texts = ax.pie(vals, labels=labels, colors=colors, startangle=90,
                           wedgeprops=dict(width=0.55, edgecolor="white", linewidth=2))
    for t in texts: t.set_fontsize(8)
    ax.text(0, 0, f"{a['pct_canalizacion']}%", ha="center", va="center",
            fontsize=16, fontweight="bold", color=C_ROJO)
    ax.set_title("Índice de Canalización", fontsize=11, fontweight="bold", pad=10)
    fig.tight_layout()
    return _grafica_bytes(fig)

def _graf_prevalencia(a: dict) -> io.BytesIO:
    """Gráfica de barras horizontales — Prevalencia de síntomas."""
    cats  = ["Recuerdos persistentes\n(Sección II)",
             "Esfuerzo por evitar\n(Sección III)",
             "Afectación\n(Sección IV)"]
    vals  = [a["pct_s2"], a["pct_s3"], a["pct_s4"]]
    n_abs = [a["prev_s2"], a["prev_s3"], a["prev_s4"]]
    cols  = [C_AZUL, C_AMBAR, C_ROJO]
    fig, ax = plt.subplots(figsize=(6, 3.5))
    bars = ax.barh(cats[::-1], vals[::-1], color=cols[::-1],
                   edgecolor="white", height=0.5)
    for bar, val, n in zip(bars, vals[::-1], n_abs[::-1]):
        ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                f"{val}%  ({n} trab.)", va="center", fontsize=8, fontweight="bold")
    ax.set_xlabel("% de trabajadores con ATS", fontsize=9)
    ax.set_xlim(0, max(vals) * 1.4 + 5 if vals else 100)
    ax.set_title("Prevalencia de Síntomas\n(entre trabajadores con ATS)",
                 fontsize=11, fontweight="bold", pad=8)
    ax.spines[["top","right"]].set_visible(False)
    ax.tick_params(axis="y", labelsize=8)
    fig.tight_layout()
    return _grafica_bytes(fig)

def _graf_segmentacion_area(a: dict) -> io.BytesIO:
    """Gráfica de barras — % de casos por área."""
    seg = a["seg_area"][:8]          # Top 8 áreas
    if not seg:
        return None
    areas = [x[0] for x in seg]
    pcts  = [x[1]["pct"] for x in seg]
    abs_  = [x[1]["atencion"] for x in seg]
    cols  = [C_ROJO if p >= 50 else C_AMBAR if p >= 25 else C_VERDE for p in pcts]
    fig, ax = plt.subplots(figsize=(6.5, max(3, len(areas) * 0.6)))
    bars = ax.barh([textwrap.shorten(a, 28) for a in areas[::-1]],
                   pcts[::-1], color=cols[::-1], edgecolor="white", height=0.55)
    for bar, val, n in zip(bars, pcts[::-1], abs_[::-1]):
        ax.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height()/2,
                f"{val}%  (n={n})", va="center", fontsize=8)
    ax.set_xlabel("% que requiere atención clínica", fontsize=9)
    ax.set_title("Segmentación por Área — Casos que Requieren Atención",
                 fontsize=10, fontweight="bold", pad=8)
    ax.spines[["top","right"]].set_visible(False)
    # Leyenda de semáforo
    patches = [
        mpatches.Patch(color=C_ROJO,   label="≥50% — Foco rojo"),
        mpatches.Patch(color=C_AMBAR,  label="25-49% — Alerta"),
        mpatches.Patch(color=C_VERDE,  label="<25% — Seguimiento"),
    ]
    ax.legend(handles=patches, fontsize=7.5, loc="lower right")
    fig.tight_layout()
    return _grafica_bytes(fig)

def _graf_urgencia(a: dict) -> io.BytesIO:
    """Gráfica de barras — Niveles de urgencia."""
    criticos = a["criticos"]
    if not criticos:
        return None
    conteo = {"CRÍTICO": 0, "ALTO": 0}
    for c in criticos:
        conteo[c["Urgencia"]] = conteo.get(c["Urgencia"], 0) + 1
    cats = list(conteo.keys())
    vals = list(conteo.values())
    cols = [C_ROJO, C_AMBAR]
    fig, ax = plt.subplots(figsize=(4.5, 3.2))
    bars = ax.bar(cats, vals, color=cols, edgecolor="white", width=0.4)
    for bar, val in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05,
                str(val), ha="center", fontsize=12, fontweight="bold")
    ax.set_ylabel("Número de trabajadores", fontsize=9)
    ax.set_title("Urgencia de Intervención", fontsize=11, fontweight="bold", pad=8)
    ax.spines[["top","right"]].set_visible(False)
    ax.set_ylim(0, max(vals) * 1.4 + 1 if vals else 5)
    fig.tight_layout()
    return _grafica_bytes(fig)


def _hoja_analiticas(wb, df: pd.DataFrame, a: dict):
    """Crea la hoja 'Analíticas' en el workbook con las 4 secciones."""
    from openpyxl.drawing.image import Image as XLImage

    ws = wb.create_sheet("Analíticas")

    # Título principal
    ws.merge_cells("A1:H1")
    ws["A1"] = "ANALÍTICAS AVANZADAS — NOM-035-STPS-2018 GUÍA I"
    ws["A1"].font      = _font(bold=True, size=14, color=BLANCO)
    ws["A1"].fill      = _fill(VERDE)
    ws["A1"].alignment = _align("center", "center")
    ws.row_dimensions[1].height = 28

    # ── SECCIÓN 1: Índice de Canalización ────────────────────────────────
    fila = 3
    ws.merge_cells(f"A{fila}:H{fila}")
    ws[f"A{fila}"] = "1. ÍNDICE DE CANALIZACIÓN"
    ws[f"A{fila}"].font      = _font(bold=True, size=11, color=BLANCO)
    ws[f"A{fila}"].fill      = _fill(ROJO)
    ws[f"A{fila}"].alignment = _align("left", "center")

    kpis = [
        ("Total evaluados",           a["total"],            VERDE),
        ("Con ATS identificado",      a["ats_total"],        ROJO if a["ats_total"] > 0 else VERDE),
        ("Requieren atención clínica",a["aten_total"],       ROJO if a["aten_total"] > 0 else VERDE),
        ("Índice de canalización",    f"{a['pct_canalizacion']}%", ROJO if a["pct_canalizacion"] > 0 else VERDE),
    ]
    fila += 1
    for label, valor, color in kpis:
        ws[f"A{fila}"] = label
        ws[f"A{fila}"].font      = _font(bold=True, size=10)
        ws[f"A{fila}"].fill      = _fill(GRIS_HEADER)
        ws[f"A{fila}"].alignment = _align("left","center")
        ws[f"A{fila}"].border    = _border()
        ws[f"B{fila}"] = valor
        ws[f"B{fila}"].font      = _font(bold=True, size=11, color=color)
        ws[f"B{fila}"].alignment = _align("center","center")
        ws[f"B{fila}"].border    = _border()
        fila += 1

    # Gráfica de dona
    buf_can = _graf_canalizacion(a)
    if buf_can:
        buf_can.seek(0)
        img = XLImage(buf_can); img.width = 300; img.height = 255
        ws.add_image(img, f"D{fila - len(kpis) - 1}")

    fila += 2

    # ── SECCIÓN 2: Prevalencia de síntomas ───────────────────────────────
    ws.merge_cells(f"A{fila}:H{fila}")
    ws[f"A{fila}"] = "2. PREVALENCIA DE SÍNTOMAS (entre trabajadores con ATS)"
    ws[f"A{fila}"].font      = _font(bold=True, size=11, color=BLANCO)
    ws[f"A{fila}"].fill      = _fill("69A2D8")
    ws[f"A{fila}"].alignment = _align("left", "center")
    fila += 1

    prev_rows = [
        ("Recuerdos persistentes (Sección II)", a["prev_s2"], f"{a['pct_s2']}%"),
        ("Esfuerzo por evitar (Sección III)",   a["prev_s3"], f"{a['pct_s3']}%"),
        ("Afectación — sueño/irritabilidad (Sección IV)", a["prev_s4"], f"{a['pct_s4']}%"),
    ]
    for etiq, n, pct in prev_rows:
        ws[f"A{fila}"] = etiq
        ws[f"A{fila}"].font      = _font(size=10); ws[f"A{fila}"].border = _border()
        ws[f"A{fila}"].fill      = _fill(GRIS_HEADER)
        ws[f"B{fila}"] = n
        ws[f"B{fila}"].font      = _font(bold=True, size=10); ws[f"B{fila}"].border = _border()
        ws[f"B{fila}"].alignment = _align("center","center")
        ws[f"C{fila}"] = pct
        ws[f"C{fila}"].font      = _font(bold=True, size=10); ws[f"C{fila}"].border = _border()
        ws[f"C{fila}"].alignment = _align("center","center")
        fila += 1

    ws[f"A{fila}"] = f"Síntoma dominante: {a['sintoma_dominante']}"
    ws[f"A{fila}"].font = _font(bold=True, size=10, color=ROJO)
    ws.merge_cells(f"A{fila}:C{fila}")
    fila += 1

    buf_prev = _graf_prevalencia(a)
    if buf_prev:
        buf_prev.seek(0)
        img = XLImage(buf_prev); img.width = 380; img.height = 240
        ws.add_image(img, f"D{fila - 6}")

    fila += 3

    # ── SECCIÓN 3: Segmentación por área ─────────────────────────────────
    ws.merge_cells(f"A{fila}:H{fila}")
    ws[f"A{fila}"] = "3. SEGMENTACIÓN POR ÁREA — FOCOS DE ATENCIÓN"
    ws[f"A{fila}"].font      = _font(bold=True, size=11, color=BLANCO)
    ws[f"A{fila}"].fill      = _fill("C8A600")
    ws[f"A{fila}"].alignment = _align("left","center")
    fila += 1

    if a["seg_area"]:
        hdrs = ["Área", "Total evaluados", "Requieren atención", "% del área", "Prioridad"]
        for col_idx, h in enumerate(hdrs, 1):
            c = ws.cell(row=fila, column=col_idx, value=h)
            c.font = _font(bold=True, size=9, color=BLANCO); c.fill = _fill(VERDE)
            c.alignment = _align("center","center"); c.border = _border(VERDE_MID)
        fila += 1
        for area, datos in a["seg_area"]:
            prioridad = "🔴 INMEDIATA" if datos["pct"] >= 50 else \
                        "🟡 ALTA"      if datos["pct"] >= 25 else "🟢 SEGUIMIENTO"
            row_data = [area, datos["total"], datos["atencion"],
                        f"{datos['pct']}%", prioridad]
            bg = ROJO_CLARO if datos["pct"] >= 50 else "FFF8DC" if datos["pct"] >= 25 else VERDE_CLARO
            for col_idx, val in enumerate(row_data, 1):
                c = ws.cell(row=fila, column=col_idx, value=val)
                c.font = _font(size=9); c.fill = _fill(bg)
                c.alignment = _align("center","center"); c.border = _border()
            fila += 1

        buf_seg = _graf_segmentacion_area(a)
        if buf_seg:
            buf_seg.seek(0)
            img = XLImage(buf_seg); img.width = 420; img.height = 260
            ws.add_image(img, f"F{fila - len(a['seg_area']) - 2}")
    else:
        ws[f"A{fila}"] = "Sin datos de área disponibles."
        fila += 1

    fila += 2

    # ── SECCIÓN 4: Urgencia de intervención ──────────────────────────────
    ws.merge_cells(f"A{fila}:H{fila}")
    ws[f"A{fila}"] = "4. URGENCIA DE INTERVENCIÓN — CASOS CRÍTICOS"
    ws[f"A{fila}"].font      = _font(bold=True, size=11, color=BLANCO)
    ws[f"A{fila}"].fill      = _fill(ROJO)
    ws[f"A{fila}"].alignment = _align("left","center")
    fila += 1

    if a["criticos"]:
        hdrs_c = ["Folio","Nombre","Área","Puesto","S2 (Sí)","S3 (Sí)","S4 (Sí)","Urgencia"]
        for col_idx, h in enumerate(hdrs_c, 1):
            c = ws.cell(row=fila, column=col_idx, value=h)
            c.font = _font(bold=True, size=9, color=BLANCO); c.fill = _fill(VERDE)
            c.alignment = _align("center","center"); c.border = _border(VERDE_MID)
        fila += 1
        for caso in sorted(a["criticos"], key=lambda x: 0 if x["Urgencia"]=="CRÍTICO" else 1):
            bg = ROJO_CLARO if caso["Urgencia"] == "CRÍTICO" else "FFF8DC"
            row_vals = [caso["Folio"], caso["Nombre"], caso["Área"],
                        caso["Puesto"], caso["S2 (Sí)"], caso["S3 (Sí)"],
                        caso["S4 (Sí)"], caso["Urgencia"]]
            for col_idx, val in enumerate(row_vals, 1):
                c = ws.cell(row=fila, column=col_idx, value=val)
                c.font = _font(bold=(caso["Urgencia"]=="CRÍTICO"), size=9)
                c.fill = _fill(bg); c.alignment = _align("left","center")
                c.border = _border()
            fila += 1

        buf_urg = _graf_urgencia(a)
        if buf_urg:
            buf_urg.seek(0)
            img = XLImage(buf_urg); img.width = 290; img.height = 220
            ws.add_image(img, f"I{fila - len(a['criticos']) - 2}")
    else:
        ws[f"A{fila}"] = "No se identificaron casos críticos o de alta urgencia."
        ws[f"A{fila}"].font = _font(bold=True, size=10, color=VERDE)
        fila += 1

    fila += 2

    # ── Párrafo de cumplimiento legal ─────────────────────────────────────
    ws.merge_cells(f"A{fila}:H{fila}")
    ws[f"A{fila}"] = "5. PÁRRAFO EJECUTIVO DE CUMPLIMIENTO LEGAL"
    ws[f"A{fila}"].font      = _font(bold=True, size=11, color=BLANCO)
    ws[f"A{fila}"].fill      = _fill(VERDE)
    ws[f"A{fila}"].alignment = _align("left","center")
    fila += 1

    ws.merge_cells(f"A{fila}:H{fila + 4}")
    c_par = ws[f"A{fila}"]
    c_par.value     = a["parrafo_cumplimiento"]
    c_par.font      = _font(size=10)
    c_par.alignment = _align("left","top", wrap=True)
    c_par.fill      = _fill(GRIS_HEADER)
    c_par.border    = _border(VERDE)
    ws.row_dimensions[fila].height = 90

    # Anchos
    ws.column_dimensions["A"].width = 32
    ws.column_dimensions["B"].width = 14
    ws.column_dimensions["C"].width = 16
    ws.column_dimensions["D"].width = 14
    ws.column_dimensions["E"].width = 14
    ws.column_dimensions["F"].width = 14
    ws.column_dimensions["G"].width = 14
    ws.column_dimensions["H"].width = 18


def _seccion_analiticas_word(doc, a: dict):
    """Agrega la sección de Analíticas Avanzadas al documento Word."""

    doc.add_page_break()

    _par(doc, "4. ANALÍTICAS AVANZADAS", bold=True, size=13, color=VERDE,
         space_before=6, space_after=4)
    p_div = doc.add_paragraph()
    b = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6")
    bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),VERDE)
    b.append(bot); p_div._p.get_or_add_pPr().append(b)
    p_div.paragraph_format.space_after = Pt(10)

    # ── 4.1 Índice de canalización ────────────────────────────────────────
    _par(doc, "4.1 Índice de Canalización", bold=True, size=11, color=ROJO,
         space_before=6, space_after=3)
    _par(doc,
         f"Del total de {a['total']} trabajadores evaluados, {a['aten_total']} "
         f"({a['pct_canalizacion']}%) requieren ser canalizados a atención médica "
         f"conforme a los criterios de la NOM-035-STPS-2018.",
         size=10, space_after=6)
    _insertar_imagen_word(doc, _graf_canalizacion(a), width_cm=9,
                          caption="Figura: Índice de Canalización")

    # ── 4.2 Prevalencia de síntomas ───────────────────────────────────────
    _par(doc, "4.2 Prevalencia de Síntomas", bold=True, size=11, color="69A2D8",
         space_before=10, space_after=3)
    _par(doc,
         f"Entre los {a['ats_total']} trabajadores con ATS identificado, el síntoma "
         f"más prevalente es '{a['sintoma_dominante']}'. La distribución por categoría es:",
         size=10, space_after=4)

    tbl_prev = doc.add_table(rows=4, cols=3)
    tbl_prev.style = "Table Grid"
    hdrs_p = ["Categoría de síntoma", "Trabajadores afectados", "% sobre ATS"]
    datos_p = [
        ("Recuerdos persistentes (S. II)", a["prev_s2"], f"{a['pct_s2']}%"),
        ("Esfuerzo por evitar (S. III)",   a["prev_s3"], f"{a['pct_s3']}%"),
        ("Afectación (S. IV)",             a["prev_s4"], f"{a['pct_s4']}%"),
    ]
    for i, h in enumerate(hdrs_p):
        _set_cell_bg(tbl_prev.cell(0,i), VERDE)
        r = tbl_prev.cell(0,i).paragraphs[0].add_run(h)
        r.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        tbl_prev.cell(0,i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_i, (etiq, n, pct) in enumerate(datos_p, 1):
        bg = VERDE_CLARO if row_i % 2 == 0 else BLANCO
        for col_i, val in enumerate([etiq, str(n), pct]):
            _set_cell_bg(tbl_prev.cell(row_i, col_i), bg)
            r = tbl_prev.cell(row_i, col_i).paragraphs[0].add_run(val)
            r.font.size=Pt(9)
            tbl_prev.cell(row_i, col_i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    _insertar_imagen_word(doc, _graf_prevalencia(a), width_cm=13,
                          caption="Figura: Prevalencia de Síntomas por Categoría")

    # ── 4.3 Segmentación por área ─────────────────────────────────────────
    if a["seg_area"]:
        _par(doc, "4.3 Segmentación por Área", bold=True, size=11, color="C8A600",
             space_before=10, space_after=3)
        _par(doc,
             f"El área con mayor concentración de casos que requieren atención clínica "
             f"es '{a['foco_rojo_area']}'. Se recomienda priorizar la intervención "
             f"en las áreas clasificadas como 'Foco Rojo' (≥50%).",
             size=10, space_after=4)
        buf_seg = _graf_segmentacion_area(a)
        if buf_seg:
            _insertar_imagen_word(doc, buf_seg, width_cm=13,
                                  caption="Figura: Porcentaje de casos por área")

    # ── 4.4 Urgencia de intervención ──────────────────────────────────────
    _par(doc, "4.4 Urgencia de Intervención — Casos Críticos", bold=True,
         size=11, color=ROJO, space_before=10, space_after=3)

    if a["criticos"]:
        _par(doc,
             f"Se identificaron {a['n_criticos']} caso(s) CRÍTICO(S) y {a['n_altos']} "
             f"caso(s) de urgencia ALTA. Los casos críticos cumplen simultáneamente "
             f"todos los criterios clínicos (S. II, III y IV) y requieren atención inmediata.",
             size=10, space_after=4)

        cols_c = ["Folio","Nombre","Área","Urgencia"]
        tbl_c = doc.add_table(rows=1, cols=len(cols_c))
        tbl_c.style = "Table Grid"
        for i, h in enumerate(cols_c):
            _set_cell_bg(tbl_c.cell(0,i), VERDE)
            r = tbl_c.cell(0,i).paragraphs[0].add_run(h)
            r.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
            tbl_c.cell(0,i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for caso in sorted(a["criticos"], key=lambda x: 0 if x["Urgencia"]=="CRÍTICO" else 1):
            row_c = tbl_c.add_row().cells
            bg = ROJO_CLARO if caso["Urgencia"]=="CRÍTICO" else "FFF8DC"
            for i, col in enumerate(cols_c):
                _set_cell_bg(row_c[i], bg)
                r = row_c[i].paragraphs[0].add_run(str(caso.get(col,"")))
                r.font.size=Pt(9)
                if caso["Urgencia"]=="CRÍTICO": r.bold=True; r.font.color.rgb=RGBColor.from_string(ROJO)
                row_c[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        buf_urg = _graf_urgencia(a)
        if buf_urg:
            _insertar_imagen_word(doc, buf_urg, width_cm=9,
                                  caption="Figura: Niveles de urgencia de intervención")
    else:
        _par(doc, "No se identificaron casos de urgencia crítica o alta.",
             bold=True, size=10, color=VERDE, space_after=4)

    # ── 4.5 Párrafo ejecutivo ─────────────────────────────────────────────
    _par(doc, "4.5 Párrafo Ejecutivo de Cumplimiento Legal", bold=True,
         size=11, color=VERDE, space_before=10, space_after=4)

    tbl_par = doc.add_table(rows=1, cols=1)
    cell_par = tbl_par.cell(0,0)
    _set_cell_bg(cell_par, VERDE_CLARO)
    _set_cell_border(cell_par, VERDE)
    r_par = cell_par.paragraphs[0].add_run(a["parrafo_cumplimiento"])
    r_par.font.size = Pt(10)
    r_par.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
    cell_par.paragraphs[0].paragraph_format.space_before = Pt(6)
    cell_par.paragraphs[0].paragraph_format.space_after  = Pt(6)
    cell_par.paragraphs[0].paragraph_format.left_indent  = Cm(0.3)
    cell_par.paragraphs[0].paragraph_format.right_indent = Cm(0.3)

REPORTES_OK = True
_REPORTE_ERROR = ''


# ── Página ────────────────────────────────────────────────────────────────────
st.set_page_config(page_title="NOM-035 · Guía I", page_icon="🏭",
                   layout="centered", initial_sidebar_state="collapsed")

# ── Modo empleado: detectar ?cliente=XXXX en la URL ───────────────────────────
_params        = st.query_params
_MODO_EMPLEADO = "cliente" in _params
_CLIENTE_URL   = str(_params.get("cliente", "")).upper() if _MODO_EMPLEADO else ""

# ── Pantalla de carga inicial ────────────────────────────────────────────────
if "app_loaded" not in st.session_state:
    st.session_state.app_loaded = False

if not st.session_state.app_loaded:
    _load_ph = st.empty()
    _load_ph.markdown(f"""
    <div class="loading-overlay">
        <img src="https://upload.wikimedia.org/wikipedia/commons/thumb/a/ac/No_image_available.svg/1px-No_image_available.svg.png"
             style="display:none">
        <div class="loading-logo" style="font-family:Montserrat,sans-serif;font-size:1.1rem;
             font-weight:700;color:#4b694e;letter-spacing:.08em;">
            RFRANYUTTI
        </div>
        <div style="font-size:.72rem;color:#888;font-family:Montserrat,sans-serif;
             font-weight:500;margin-top:.3rem;letter-spacing:.1em;">
            CONCIENCIA VERDE Y LABORAL S.C.
        </div>
        <div class="loading-ring"></div>
        <div class="loading-txt">NOM-035-STPS-2018 · Cargando...</div>
    </div>
    """, unsafe_allow_html=True)
    import time as _t; _t.sleep(1.2)
    _load_ph.empty()
    st.session_state.app_loaded = True

# ── Clientes ──────────────────────────────────────────────────────────────────
CLIENTES = {
    "FRUCO": {
        "razon": "FRUTAS CONCENTRADAS, S.A.P.I. DE C.V.",
        "logo":  "assets/logos/fruco.png",        # PNG sin fondo (remove_bg.py)
        "opciones": ["FRUTAS CONCENTRADAS, S.A.P.I. DE C.V."],
    },
    "QUALTIA": {
        "razon": "QUALTIA ALIMENTOS Y OPERACIONES, S. DE R.L. DE C.V.",
        "logo":  "assets/logos/qualtia.png",       # convertido a PNG sin fondo
        "opciones": [
            "QUALTIA ALIMENTOS Y OPERACIONES, S. DE R.L. DE C.V.",
            "QUALTIA ALIMENTOS OPERACIONES, S. DE R.L. DE C.V. (CEDIS Y SERVICIOS AUXILIARES)",
        ],
    },
    "DIABLOS": {
        "razon": "CENTRO DEPORTIVO ALFREDO HARP HELÚ, S.A. DE C.V.",
        "logo":  "assets/logos/Diablos.png",       # convertido a PNG sin fondo
        "opciones": ["CENTRO DEPORTIVO ALFREDO HARP HELÚ, S.A. DE C.V."],
    },
}
LOGO_RF = "assets/logos/rfranyutti.gif"

def excel_path(cliente_key: str, razon_social: str = "") -> str:
    """
    Un archivo Excel por cliente.
    QUALTIA genera dos archivos según razón social:
      · resultados_QUALTIA.xlsx       → planta principal
      · resultados_QUALTIA_CEDIS.xlsx → CEDIS Y SERVICIOS AUXILIARES
    """
    if cliente_key == "QUALTIA" and "CEDIS" in razon_social.upper():
        return "data/resultados_QUALTIA_CEDIS.xlsx"
    return f"data/resultados_{cliente_key.upper()}.xlsx"

# ── Catálogos ──────────────────────────────────────────────────────────────────
SEL = "— Selecciona —"

OPC_SEXO  = [SEL, "Femenino", "Masculino"]
OPC_EDAD  = [SEL,"15 - 19","20 - 24","25 - 29","30 - 34",
             "35 - 39","40 - 44","45 - 49","50 - 54",
             "55 - 59","60 - 64","65 - 69","70 o más"]
OPC_ECIVIL = [SEL,"Soltero","Casado","Unión libre","Divorciado","Viudo"]
OPC_ESTUD  = [SEL,"Sin formación","Primaria","Secundaria",
              "Preparatoria o Bachillerato","Técnico Superior",
              "Licenciatura","Maestría","Doctorado"]
OPC_PUESTO = [SEL,"Operativo","Supervisor","Profesional o técnico","Gerente"]
OPC_CONTRAT = [SEL,"Por obra o proyecto","Tiempo indeterminado",
               "Por tiempo determinado (temporal)","Honorarios"]
OPC_PERSONAL = [SEL,"Sindicalizado","Confianza","Ninguno"]
OPC_JORNADA  = [SEL,
                "Fijo diurno (entre las 6:00 y 20:00 hrs)",
                "Fijo nocturno (entre las 20:00 y 6:00 hrs)",
                "Fijo mixto (combinación de nocturno y diurno)"]
OPC_TPUESTO = [SEL,"Menos de 6 meses","Entre 6 meses y 1 año",
               "Entre 1 a 4 años","Entre 5 a 9 años",
               "Entre 10 a 14 años","Entre 15 a 19 años",
               "Entre 20 a 24 años","25 años o más"]
OPC_EXP     = [SEL,"Menos de 6 meses","Entre 6 meses y 1 año",
               "Entre 1 a 4 años","Entre 5 a 9 años",
               "Entre 10 a 14 años","Entre 15 a 19 años",
               "Entre 20 a 24 años","25 años o más"]

# ── Preguntas Guía I (texto oficial NOM-035) ───────────────────────────────────
S1 = [
    "Accidente que tenga como consecuencia la muerte, la pérdida de un miembro o una lesión grave.",
    "Asaltos.",
    "Actos violentos que derivaron en lesiones graves.",
    "Secuestro.",
    "Amenazas.",
    "Cualquier otro acontecimiento que ponga en riesgo su vida o salud, y/o la de otras personas.",
]
S2 = [
    "¿Ha tenido recuerdos recurrentes sobre el acontecimiento que le provocan malestares?",
    "¿Ha tenido sueños de carácter recurrente sobre el acontecimiento, que le producen malestar?",
]
S3 = [
    "¿Se ha esforzado por evitar todo tipo de sentimientos, conversaciones o situaciones que le puedan recordar el acontecimiento?",
    "¿Se ha esforzado por evitar todo tipo de actividades, lugares o personas que motivan recuerdos del acontecimiento?",
    "¿Ha tenido dificultad para recordar alguna parte importante del evento?",
    "¿Ha disminuido su interés en sus actividades cotidianas?",
    "¿Se ha sentido usted alejado o distante de los demás?",
    "¿Ha notado que tiene dificultad para expresar sus sentimientos?",
    "¿Ha tenido la impresión de que su vida se va a acortar, que va a morir antes que otras personas o que tiene un futuro limitado?",
]
S4 = [
    "¿Ha tenido usted dificultades para dormir?",
    "¿Ha estado particularmente irritable o le han dado arranques de coraje?",
    "¿Ha tenido dificultad para concentrarse?",
    "¿Ha estado nervioso o constantemente en alerta?",
    "¿Se ha sobresaltado fácilmente por cualquier cosa?",
]

SECCIONES = [
    {"id":"s1","titulo":"I. Acontecimiento traumático severo",
     "inst":"¿Ha presenciado o sufrido alguna vez, durante o con motivo del trabajo:","preg":S1},
    {"id":"s2","titulo":"II. Recuerdos persistentes sobre el acontecimiento",
     "inst":"(durante el último mes):","preg":S2},
    {"id":"s3","titulo":"III. Esfuerzo por evitar circunstancias parecidas o asociadas al acontecimiento",
     "inst":"(durante el último mes):","preg":S3},
    {"id":"s4","titulo":"IV. Afectación",
     "inst":"(durante el último mes):","preg":S4},
]

# ── Lógica resultado NOM-035 ───────────────────────────────────────────────────
def evaluar(r1, r2, r3, r4):
    pos1 = sum(1 for r in r1 if r == "Sí")
    if pos1 == 0:
        return dict(ats=False, atencion=False,
                    nivel="NO REQUIERE VALORACIÓN CLÍNICA",
                    desc="Todas las respuestas a la Sección I son 'No'. "
                         "El trabajador no requiere valoración clínica.",
                    criterios="Ninguno",
                    p1=0,p2=0,p3=0,p4=0)

    pos2 = sum(1 for r in r2 if r == "Sí")
    pos3 = sum(1 for r in r3 if r == "Sí")
    pos4 = sum(1 for r in r4 if r == "Sí")
    crit = []
    if pos2 >= 1: crit.append("Sección II: al menos 1 respuesta 'Sí' en recuerdos persistentes")
    if pos3 >= 3: crit.append(f"Sección III: {pos3} respuestas 'Sí' en esfuerzo por evitar (criterio ≥3)")
    if pos4 >= 2: crit.append(f"Sección IV: {pos4} respuestas 'Sí' en afectación (criterio ≥2)")
    req = len(crit) > 0
    return dict(
        ats=True, atencion=req,
        nivel="REQUIERE ATENCIÓN CLÍNICA" if req else "IDENTIFICADO — SIN CRITERIO DE ATENCIÓN CLÍNICA",
        desc=("Expuesto a ATS en el trabajo. "
              + ("Criterios activados: " + "; ".join(crit) if req
                 else "No se activan criterios de atención clínica en Secciones II, III y IV.")),
        criterios="; ".join(crit) if crit else "Ninguno",
        p1=pos1,p2=pos2,p3=pos3,p4=pos4)

# ── Excel ──────────────────────────────────────────────────────────────────────
def init_excel(path: str):
    os.makedirs("data", exist_ok=True)
    if not os.path.exists(path):
        cols = ["Folio","Fecha","Cliente","Razón Social","Nombre","Sexo","Edad",
                "Estado Civil","Nivel Estudios","Estud. Status","Puesto","Área",
                "Contratación","Tipo Personal","Jornada","Rotación Turnos",
                "Tiempo Puesto","Experiencia","ATS","Atención Clínica",
                "Nivel","Descripción","Criterios","P-S1","P-S2","P-S3","P-S4"]
        for i in range(1,7):  cols.append(f"S1P{i:02d}")
        for i in range(1,3):  cols.append(f"S2P{i:02d}")
        for i in range(1,8):  cols.append(f"S3P{i:02d}")
        for i in range(1,6):  cols.append(f"S4P{i:02d}")
        pd.DataFrame(columns=cols).to_excel(path, index=False)

def guardar(data):
    path = excel_path(data["cliente"], data.get("razon",""))
    init_excel(path)
    for intento in range(5):
        try:
            df  = pd.read_excel(path)
            res = data["res"]
            fila = {
                "Folio":data["folio"],"Fecha":datetime.now().strftime("%Y-%m-%d %H:%M"),
                "Cliente":data["cliente"],"Razón Social":data["razon"],
                "Nombre":data["nombre"],"Sexo":data["sexo"],"Edad":data["edad"],
                "Estado Civil":data["ecivil"],"Nivel Estudios":data["estudios"],
                "Estud. Status":data["estatus"],"Puesto":data["puesto"],"Área":data["area"],
                "Contratación":data["contrat"],"Tipo Personal":data["personal"],
                "Jornada":data["jornada"],"Rotación Turnos":data["rotacion"],
                "Tiempo Puesto":data["tpuesto"],"Experiencia":data["exp"],
                "ATS":"Sí" if res["ats"] else "No",
                "Atención Clínica":"Sí" if res["atencion"] else "No",
                "Nivel":res["nivel"],"Descripción":res["desc"],
                "Criterios":res["criterios"],
                "P-S1":res["p1"],"P-S2":res["p2"],"P-S3":res["p3"],"P-S4":res["p4"],
            }
            for i,r in enumerate(data.get("r1",[]),1): fila[f"S1P{i:02d}"] = r
            for i,r in enumerate(data.get("r2",[]),1): fila[f"S2P{i:02d}"] = r
            for i,r in enumerate(data.get("r3",[]),1): fila[f"S3P{i:02d}"] = r
            for i,r in enumerate(data.get("r4",[]),1): fila[f"S4P{i:02d}"] = r
            df = pd.concat([df, pd.DataFrame([fila])], ignore_index=True)
            df.to_excel(path, index=False)
            return True
        except PermissionError:
            if intento < 4: time.sleep(1)
            else:
                st.error("⚠ El archivo Excel está abierto. Ciérralo y presiona SIGUIENTE nuevamente.")
                return False
    return False

# Mutex en memoria — garantiza orden en procesos Streamlit multi-hilo
_folio_lock = threading.Lock()

def _db_path(cliente_key: str) -> str:
    os.makedirs('data', exist_ok=True)
    return f"data/folios_{cliente_key.upper()}.db"

def _init_db(db_path: str):
    with sqlite3.connect(db_path, timeout=30) as con:
        con.execute('''CREATE TABLE IF NOT EXISTS folios (
            id    INTEGER PRIMARY KEY AUTOINCREMENT,
            razon TEXT NOT NULL,
            ts    TEXT NOT NULL
        )''')
        con.commit()

def folio_nuevo(cliente_key: str, razon_social: str) -> str:
    """Folio 001, 002, 003... estrictamente único por orden de llegada.
    SQLite AUTOINCREMENT + threading.Lock garantizan que nunca
    dos usuarios obtengan el mismo número aunque entren simultáneamente.
    """
    db        = _db_path(cliente_key)
    _init_db(db)
    razon_key = razon_social.strip().upper()
    with _folio_lock:
        with sqlite3.connect(db, timeout=30, check_same_thread=False) as con:
            cur = con.execute(
                'INSERT INTO folios (razon, ts) VALUES (?, ?)',
                (razon_key, datetime.now().isoformat())
            )
            con.commit()
            n = con.execute(
                'SELECT COUNT(*) FROM folios WHERE razon=? AND id<=?',
                (razon_key, cur.lastrowid)
            ).fetchone()[0]
    return str(n).zfill(3)

def solo_letras(t):
    return re.sub(r"[^A-Za-záéíóúÁÉÍÓÚüÜñÑ\s;]","",t).upper().strip()

def nombre_tiene_caracteres_invalidos(t: str) -> bool:
    """True si el texto contiene números o símbolos no permitidos."""
    return bool(re.search(r"[^A-Za-záéíóúÁÉÍÓÚüÜñÑ\s]", t))

def trabajador_ya_registrado(ap1: str, ap2: str, nom: str, razon: str, cliente_key: str) -> bool:
    """Verifica si ya existe un registro con el mismo nombre completo en la misma razón social."""
    path = excel_path(cliente_key, razon)
    if not os.path.exists(path):
        return False
    try:
        df = pd.read_excel(path)
        if df.empty or "Nombre" not in df.columns:
            return False
        nombre_nuevo = f"{ap1}; {ap2}; {nom}".strip().upper()
        mask = (
            (df["Nombre"].str.strip().str.upper() == nombre_nuevo) &
            (df["Razón Social"].str.strip().str.upper() == razon.strip().upper())
        )
        return mask.any()
    except Exception:
        return False

def idx_de(lst, val):
    return lst.index(val) if val in lst else 0

# ── CSS ────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Montserrat:ital,wght@0,300;0,400;0,500;0,600;0,700;1,400&display=swap');
:root{--v:#4b694e;--v2:#6a9370;--az:#69a2d8;--rj:#a20000;--gr:#b4b4b4;
      --fo:#f1f1f1;--bl:#ffffff;--tx:#1a1a1a;--br:#d0d0d0;--am:#c8a600;}
html,body,.stApp{background:var(--fo)!important;font-family:'Montserrat',sans-serif!important;color:var(--tx);}
#MainMenu,footer,header{visibility:hidden;}
.block-container{padding:1.6rem 2rem!important;max-width:860px!important;margin:auto;}

/* ── Folio ── */
.folio-box{font-size:.78rem;font-weight:600;color:#888;text-align:right;letter-spacing:.06em;text-transform:uppercase;line-height:1.4;}
.folio-num{font-size:1.25rem;font-weight:700;color:var(--tx);}

/* ── Panel ── */
.ptitulo{font-size:1.35rem;font-weight:700;color:var(--v);border-bottom:2px solid var(--v);padding-bottom:.4rem;margin-bottom:1.2rem;}
.pcard{background:var(--bl);border-radius:13px;padding:1.4rem 1.8rem;margin-bottom:1rem;box-shadow:0 2px 8px rgba(0,0,0,.07);}

/* ── Bienvenida ── */
.bv-box{background:var(--bl);border-radius:15px;padding:2rem 2.5rem;text-align:center;box-shadow:0 4px 14px rgba(0,0,0,.09);margin-bottom:1.4rem;}
.bv-tit{font-size:1.55rem;font-weight:700;color:var(--v);text-transform:uppercase;letter-spacing:.05em;margin-bottom:.3rem;}
.bv-sub{font-size:.88rem;font-weight:500;color:#666;text-transform:uppercase;letter-spacing:.1em;margin-bottom:1.2rem;}
.priv{background:#f8f8f8;border:1px solid var(--br);border-radius:9px;padding:1rem 1.4rem;text-align:left;font-size:.92rem;color:#555;margin-bottom:1.2rem;line-height:1.8;}

/* ── Labels ── */
.slabel{font-size:.78rem;font-weight:700;color:var(--v);text-transform:uppercase;letter-spacing:.1em;margin:.9rem 0 .25rem 0;}

/* ── Confirmación / Aviso ── */
.cfm-box{background:var(--bl);border:2px solid var(--v);border-radius:13px;padding:2rem 2.4rem;text-align:center;font-size:1rem;line-height:1.9;color:var(--tx);margin-bottom:1.4rem;box-shadow:0 4px 14px rgba(75,105,78,.14);}
.av-box{background:var(--bl);border:2px solid var(--am);border-radius:13px;padding:1.7rem 2.2rem;text-align:center;font-size:1rem;color:var(--tx);line-height:1.9;margin-bottom:1.4rem;box-shadow:0 4px 14px rgba(200,166,0,.11);}
.av-tit{font-size:1.05rem;font-weight:700;color:#7a5900;margin-bottom:.5rem;}

/* ── Pregunta ── */
.pq-card{background:var(--bl);border-radius:13px;padding:1.7rem 2.1rem;margin-bottom:.85rem;box-shadow:0 3px 11px rgba(0,0,0,.08);}
.pq-sec{font-size:.75rem;font-weight:700;color:var(--v);text-transform:uppercase;letter-spacing:.1em;margin-bottom:.2rem;}
.pq-ins{font-size:.85rem;color:#888;margin-bottom:.35rem;}
.pq-txt{font-size:1.08rem;font-weight:600;color:var(--tx);line-height:1.6;}
.err-r{color:var(--rj);font-size:.88rem;font-weight:700;margin-top:.4rem;}
.prog-txt{font-size:.82rem;color:#888;text-align:right;margin-bottom:.5rem;font-weight:500;}

/* ── Resultado ── */
.res-v{background:var(--v);color:#fff;border-radius:13px;padding:1.9rem 2.1rem;text-align:center;margin-bottom:1rem;}
.res-r{background:var(--rj);color:#fff;border-radius:13px;padding:1.9rem 2.1rem;text-align:center;margin-bottom:1rem;}
.res-niv{font-size:1.25rem;font-weight:700;margin-top:.4rem;}
.res-sub{font-size:.86rem;opacity:.8;margin-top:.4rem;}
.ac-box{background:#fffbe6;border:1.5px solid var(--am);border-radius:9px;padding:1rem 1.4rem;font-size:.93rem;color:#5a4a00;margin-bottom:1rem;}
.dt-box{background:#f3f3f3;border-radius:9px;padding:.85rem 1.2rem;font-size:.88rem;color:#444;margin-bottom:.6rem;}

/* ── Fin ── */
.fin-box{background:var(--v);color:#fff;border-radius:17px;padding:3rem 2rem;text-align:center;box-shadow:0 6px 22px rgba(75,105,78,.24);}
.fin-tit{font-size:1.35rem;font-weight:700;letter-spacing:.05em;margin-bottom:.6rem;}
.fin-sub{font-size:.96rem;opacity:.84;}

/* ── Botones ── */
.stButton>button{font-family:'Montserrat',sans-serif!important;font-weight:700!important;font-size:.93rem!important;
  letter-spacing:.05em!important;border-radius:9px!important;border:none!important;
  padding:.65rem 1.8rem!important;cursor:pointer!important;
  background:var(--v)!important;color:#fff!important;transition:opacity .2s!important;}
.stButton>button:hover{opacity:.84!important;}
div[role="radiogroup"]>label{background:#fafafa;border:1.5px solid var(--br);border-radius:7px;
  padding:.5rem 1rem!important;margin-bottom:.35rem!important;font-size:.96rem!important;
  font-weight:500;cursor:pointer;transition:border-color .14s,background .14s;}
div[role="radiogroup"]>label:hover{border-color:var(--v2);background:#f0f6f0;}
hr.div{border:none;border-top:1.5px solid var(--br);margin:.8rem 0;}

/* ── Animación de carga ── */
@keyframes rf-pulse{
  0%,100%{opacity:1;transform:scale(1);}
  50%{opacity:.55;transform:scale(.96);}
}
@keyframes rf-spin{
  0%{transform:rotate(0deg);}100%{transform:rotate(360deg);}
}
.loading-overlay{
  position:fixed;inset:0;background:rgba(241,241,241,.92);
  display:flex;flex-direction:column;align-items:center;justify-content:center;
  z-index:9999;
}
.loading-logo{animation:rf-pulse 1.4s ease-in-out infinite;}
.loading-ring{
  width:48px;height:48px;border:4px solid #d0d0d0;
  border-top-color:var(--v);border-radius:50%;
  animation:rf-spin .9s linear infinite;margin-top:1.2rem;
}
.loading-txt{font-size:.82rem;font-weight:600;color:#4b694e;margin-top:.7rem;letter-spacing:.06em;}

/* campos nombre */
.stTextInput input {
    text-transform: uppercase !important;
    font-family: Montserrat, sans-serif !important;
    font-weight: 600 !important;
    letter-spacing: .04em !important;
}
/* alerta duplicado */
.dup-alert {
    background: #fff0f0;
    border: 2px solid #a20000;
    border-radius: 11px;
    padding: 1rem 1.4rem;
    margin: .8rem 0;
    font-size: .88rem;
    color: #a20000;
    font-weight: 600;
    line-height: 1.7;
}
/* alerta campo inválido */
.campo-err {
    color: #a20000;
    font-size: .75rem;
    font-weight: 700;
    margin-top: .1rem;
    display: block;
}
</style>
""", unsafe_allow_html=True)

# Validación de nombre: se realiza en Python al presionar INICIAR CUESTIONARIO

# ── Estado ─────────────────────────────────────────────────────────────────────
# Modo empleado: cliente y pantalla inicial desde URL
_cliente_def  = _CLIENTE_URL if (_MODO_EMPLEADO and _CLIENTE_URL in CLIENTES) else "FRUCO"
_pantalla_def = "bienvenida" if _MODO_EMPLEADO else "panel"

DEF = dict(
    pantalla=_pantalla_def, cliente_key=_cliente_def,
    razon=CLIENTES[_cliente_def]["opciones"][0],
    areas=["Producción","Administración","Recursos Humanos","Ventas","Operaciones"],
    folio="001",
    ap1="", ap2="", nom="",
    sexo=SEL, edad=SEL, ecivil=SEL,
    estudios=SEL, estatus="Terminada",
    puesto=SEL, area=SEL,
    contrat=SEL, personal=SEL,
    jornada=SEL, rotacion="Sí",
    tpuesto=SEL, exp=SEL,
    sec=0, preg=0,
    r1=[], r2=[], r3=[], r4=[],
    skip=False, err=False, res=None, modal=None,
    form_v=0,
)
for k,v in DEF.items():
    if k not in st.session_state: st.session_state[k] = v
S = st.session_state

# Si modo empleado pero la sesión tiene pantalla=panel (sesión reutilizada),
# forzar a bienvenida para que el empleado no vea el panel
if _MODO_EMPLEADO and S.get("pantalla") == "panel":
    S["pantalla"] = "bienvenida"
    S["cliente_key"] = _cliente_def
    S["razon"] = CLIENTES[_cliente_def]["opciones"][0]

def limpiar():
    """Resetea las variables de estado del formulario."""
    for k in ["ap1","ap2","nom"]: S[k]=""
    for k in ["sexo","edad","ecivil","estudios","puesto","area",
              "contrat","personal","jornada","tpuesto","exp"]: S[k]=SEL
    S.estatus="Terminada"; S.rotacion="Sí"

def borrar_formulario():
    """Reset completo del formulario.
    Incrementar form_v cambia todas las keys de widgets — Streamlit los
    renderiza desde cero, vacíos, sin importar lo que el usuario ingresó."""
    limpiar()
    S.form_v = S.get("form_v", 0) + 1

def reset_cuestionario():
    S.sec=0; S.preg=0
    S.r1=[]; S.r2=[]; S.r3=[]; S.r4=[]
    S.skip=False; S.err=False; S.res=None

def _img_b64(path: str) -> str:
    """Convierte imagen a base64 para incrustarla en HTML sin depender de st.image."""
    import base64, mimetypes
    try:
        mime = mimetypes.guess_type(path)[0] or "image/png"
        with open(path, "rb") as f:
            return f"data:{mime};base64,{base64.b64encode(f.read()).decode()}"
    except Exception:
        return ""

def header(folio=False):
    """Header HTML con ambos logos alineados al centro, mismo alto visual."""
    rf_src     = _img_b64(LOGO_RF)
    cli_src    = _img_b64(CLIENTES[S.cliente_key]["logo"])
    folio_html = (
        f'<div class="folio-box" style="margin-left:auto;text-align:right;">'
        f'NO. DE CUESTIONARIO<br>'
        f'<span class="folio-num">{S.folio}</span></div>'
    ) if folio else ""

    rf_tag  = (f'<img src="{rf_src}"  alt="RFRANYUTTI" ' +
               'style="height:52px;max-width:120px;width:auto;object-fit:contain;">') if rf_src else "<b>RFRANYUTTI</b>"
    cli_tag = (f'<img src="{cli_src}" alt="{S.cliente_key}" ' +
               'style="height:52px;max-width:120px;width:auto;object-fit:contain;">') if cli_src else f"<b>{S.cliente_key}</b>"

    st.markdown(f"""
    <div style="display:flex;flex-wrap:wrap;align-items:center;
                justify-content:space-between;gap:.5rem;
                padding:.5rem 0 .6rem 0;margin-bottom:.2rem;">
        <div style="display:flex;align-items:center;gap:.8rem;flex-wrap:wrap;">
            {rf_tag}
            <div style="width:1px;height:44px;background:#ccc;flex-shrink:0;"></div>
            {cli_tag}
        </div>
        <div style="flex-shrink:0;">{folio_html}</div>
    </div>
    <hr class="div">
    """, unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# PANEL DEL OPERATIVO
# ══════════════════════════════════════════════════════════════════════════════
if S.pantalla == "panel":
    st.markdown('<div class="ptitulo">⚙ PANEL DEL OPERATIVO · NOM-035-STPS-2018</div>',
                unsafe_allow_html=True)

    # ── DIAGNÓSTICO ────────────────────────────────────────────────────────
    import glob as _glob
    _cwd = os.getcwd()
    _candidatos = [
        os.path.join(_cwd, "generar_reporte.py"),
        os.path.join(os.path.abspath("."), "generar_reporte.py"),
    ]
    try:
        _candidatos.append(os.path.join(os.path.dirname(os.path.abspath(__file__)), "generar_reporte.py"))
    except: pass

    _encontrado = [c for c in _candidatos if os.path.exists(c)]
    _pys = [os.path.basename(p) for p in _glob.glob(os.path.join(_cwd, "*.py"))]

    with st.expander("🔍 Diagnóstico de reportes" + (" — ✅ OK" if REPORTES_OK else " — ❌ ERROR"), expanded=not REPORTES_OK):
        st.write(f"**REPORTES_OK:** `{REPORTES_OK}`")
        st.write(f"**Carpeta actual:** `{_cwd}`")
        st.write(f"**Archivos .py encontrados:** `{_pys}`")
        for c in _candidatos:
            st.write(f"{'✅' if os.path.exists(c) else '❌'} `{c}`")
        if _REPORTE_ERROR:
            st.code(_REPORTE_ERROR, language="text")
        if not REPORTES_OK and _encontrado:
            st.warning("⚠ El archivo existe pero no se puede importar. Ver error arriba.")
            if st.button("🔄 Reintentar importación"):
                import importlib, sys as _sys
                if "generar_reporte" in _sys.modules:
                    del _sys.modules["generar_reporte"]
                _sys.path.insert(0, os.path.dirname(_encontrado[0]))
                try:
                    from generar_reporte import generar_excel_mejorado, generar_word
                    st.success("✅ Importación exitosa. Recarga la página.")
                except Exception as _e2:
                    st.error(f"Error: {_e2}")

    if S.modal == "borrar":
        st.warning("⚠ **LOS REGISTROS SE ELIMINARÁN Y SE BORRARÁN TODOS LOS DATOS CONTENIDOS.**\n\n¿Desea continuar?")
        c1,c2 = st.columns(2)
        with c1:
            if st.button("CONTINUAR — BORRAR TODO"):
                p = excel_path(S.cliente_key, S.razon)
                if os.path.exists(p): os.remove(p)
                S.modal=None; st.success("✓ Registros eliminados."); st.rerun()
        with c2:
            if st.button("CANCELAR"): S.modal=None; st.rerun()
        st.stop()

    if S.modal == "terminar":
        st.info("ℹ **LOS DATOS SE REGISTRARÁN EN EL SISTEMA Y NO SERÁ POSIBLE REALIZAR CAMBIOS POSTERIORES NI VOLVER A COMPLETAR EL CUESTIONARIO.**\n\n¿Desea continuar?")
        c1,c2 = st.columns(2)
        with c1:
            if st.button("CONTINUAR — TERMINAR REGISTROS"):
                S.modal=None; st.success("✓ Sesión cerrada."); st.rerun()
        with c2:
            if st.button("CANCELAR"): S.modal=None; st.rerun()
        st.stop()

    st.markdown('<div class="pcard">', unsafe_allow_html=True)
    st.markdown("**1. Selección de cliente**")
    ck = st.selectbox("Cliente", list(CLIENTES.keys()),
                      index=list(CLIENTES.keys()).index(S.cliente_key))
    S.cliente_key = ck
    inf = CLIENTES[ck]
    S.razon = st.selectbox("Razón social", inf["opciones"])
    try: st.image(inf["logo"], width=120)
    except: st.caption(f"Logo no encontrado: {inf['logo']}")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="pcard">', unsafe_allow_html=True)
    st.markdown("**2. Departamentos / Áreas**")
    st.caption("Una área por línea. Se mostrarán como opciones en el cuestionario.")
    at = st.text_area("", "\n".join(S.areas), height=120, label_visibility="collapsed")
    S.areas = [a.strip() for a in at.split("\n") if a.strip()]
    st.caption(f"{len(S.areas)} área(s) configuradas.")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="pcard">', unsafe_allow_html=True)
    st.markdown("**3. Guía a aplicar**")
    st.info("**GUÍA I** — Identificación de Trabajadores Sujetos a "
            "Acontecimientos Traumáticos Severos")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("**4. Acciones**")
    c1,c2,c3,c4 = st.columns(4)
    with c1:
        _p_raw = excel_path(S.cliente_key, S.razon)
        if os.path.exists(_p_raw):
            with open(_p_raw, "rb") as _f:
                st.download_button(
                    label="⬇️  DESCARGAR DATOS (Excel)",
                    data=_f.read(),
                    file_name=os.path.basename(_p_raw),
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
        else:
            st.button("⬇️  DESCARGAR DATOS (Excel)", disabled=True,
                      use_container_width=True,
                      help=f"Sin datos aún para {S.cliente_key}.")
    with c2:
        if st.button("ABRIR WORD"):
            p = excel_path(S.cliente_key, S.razon)
            if REPORTES_OK and os.path.exists(p):
                with st.spinner("Generando Word..."):
                    out = generar_word(p, S.cliente_key, S.razon,
                                       logo_rf=LOGO_RF,
                                       logo_cliente=CLIENTES[S.cliente_key]["logo"])
                if out: st.success(f"Word generado: `{os.path.abspath(out)}`")
                else:   st.warning("Sin datos suficientes.")
            elif not REPORTES_OK: st.error("generar_reporte.py no encontrado.")
            else: st.warning(f"Sin datos aún para {S.cliente_key}.")
    with c3:
        if st.button("BORRAR REGISTROS"): S.modal="borrar"; st.rerun()
    with c4:
        if st.button("TERMINAR REGISTROS"): S.modal="terminar"; st.rerun()

    st.markdown('<hr class="div">', unsafe_allow_html=True)
    st.markdown("**5. Generar Reportes**")

    p_rep = excel_path(S.cliente_key, S.razon)
    _hay_datos = os.path.exists(p_rep)
    col_r1, col_r2 = st.columns(2)
    with col_r1:
        if st.button("📊  EXCEL MEJORADO + GRÁFICAS", use_container_width=True,
                     disabled=not _hay_datos):
            if not REPORTES_OK:
                st.error(f"⚠ No se encontró generar_reporte.py — Error: {_REPORTE_ERROR}")
            else:
                with st.spinner("Generando Excel con formatos y gráficas..."):
                    try:
                        out = generar_excel_mejorado(p_rep, S.cliente_key, S.razon)
                        if out:
                            with open(out, "rb") as f:
                                _excel_bytes = f.read()
                            _nombre_excel = os.path.basename(out)
                            st.success("✓ Excel listo — haz click para descargar:")
                            st.download_button(
                                label="⬇️  DESCARGAR EXCEL",
                                data=_excel_bytes,
                                file_name=_nombre_excel,
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                use_container_width=True,
                            )
                        else:
                            st.warning("Sin datos suficientes para generar el reporte.")
                    except Exception as e:
                        st.error(f"Error al generar Excel: {e}")
    with col_r2:
        if st.button("📄  INFORME WORD COMPLETO", use_container_width=True,
                     disabled=not _hay_datos):
            if not REPORTES_OK:
                st.error(f"⚠ No se encontró generar_reporte.py — Error: {_REPORTE_ERROR}")
            else:
                with st.spinner("Generando informe Word con portada y gráficas..."):
                    try:
                        out = generar_word(p_rep, S.cliente_key, S.razon,
                                           logo_rf=LOGO_RF,
                                           logo_cliente=CLIENTES[S.cliente_key]["logo"])
                        if out:
                            with open(out, "rb") as f:
                                _word_bytes = f.read()
                            _nombre_word = os.path.basename(out)
                            st.success("✓ Word listo — haz click para descargar:")
                            st.download_button(
                                label="⬇️  DESCARGAR WORD",
                                data=_word_bytes,
                                file_name=_nombre_word,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                            )
                        else:
                            st.warning("Sin datos suficientes para generar el reporte.")
                    except Exception as e:
                        st.error(f"Error al generar Word: {e}")
    if not _hay_datos:
        st.caption("⚠ Registra al menos un cuestionario antes de generar reportes.")
    if not REPORTES_OK:
        st.error(f"⚠ generar_reporte.py no encontrado.\n"
                 f"Carpeta actual: `{os.getcwd()}`\n"
                 f"Coloca generar_reporte.py en esa carpeta y reinicia la app.")

    st.markdown('<hr class="div">', unsafe_allow_html=True)
    st.markdown("**6. Link para empleados**")
    _base = st.text_input(
        "URL base de tu app",
        value="http://localhost:8501",
        help="Cuando actives ngrok pega aquí tu URL. Ejemplo: https://abc123.ngrok-free.app"
    )
    _link = f"{_base.rstrip('/')}/?cliente={S.cliente_key}"
    st.markdown(f"""
    <div style="background:#f0f6f0;border:2px solid #4b694e;border-radius:10px;
                padding:1rem 1.4rem;margin:.4rem 0 .8rem 0;font-family:Montserrat,sans-serif;">
        <div style="font-size:.72rem;font-weight:700;color:#4b694e;text-transform:uppercase;
                    letter-spacing:.08em;margin-bottom:.35rem;">
            🔗 Link empleados — {S.cliente_key}
        </div>
        <div style="font-size:.9rem;font-weight:600;word-break:break-all;background:#fff;
                    border-radius:6px;padding:.45rem .8rem;border:1px solid #d0d0d0;">
            {_link}
        </div>
        <div style="font-size:.75rem;color:#666;margin-top:.4rem;">
            Comparte por WhatsApp · Sesiones independientes · Folio automático por empleado
        </div>
    </div>
    """, unsafe_allow_html=True)
    st.code(_link, language=None)

    st.markdown('<hr class="div">', unsafe_allow_html=True)
    st.markdown("**7. Cuestionario modo local**")
    if st.button("🟢  INICIAR CUESTIONARIO PARA EMPLEADO", use_container_width=True):
        S.folio=folio_nuevo(S.cliente_key, S.razon); limpiar(); reset_cuestionario()
        S.pantalla="bienvenida"; st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# BIENVENIDA
# ══════════════════════════════════════════════════════════════════════════════
elif S.pantalla == "bienvenida":
    header(False)
    st.markdown("""
    <div class="bv-box">
      <div class="bv-tit">NOM-035-STPS-2018</div>
      <div class="bv-sub">Cuestionario para identificar a los trabajadores<br>
        que fueron sujetos a acontecimientos traumáticos severos</div>
      <p style="font-size:.83rem;color:#444;line-height:1.8;text-align:left;">
        Este cuestionario tiene el propósito de identificar y medir las condiciones del entorno
        psicosocial que pueden afectar la salud y el bienestar de las personas trabajadoras.
        Sus respuestas son fundamentales para construir un diagnóstico confiable y orientar
        acciones de mejora. Le pedimos responder con honestidad, con base en su experiencia
        en el centro de trabajo.
      </p>
      <div class="bv-sub" style="margin-top:.9rem;">Cláusula de Privacidad</div>
      <div class="priv">
        La información proporcionada será tratada con estricta confidencialidad. Los datos y
        resultados se utilizarán exclusivamente para fines internos de diagnóstico y mejora del
        ambiente laboral, conforme a la normativa vigente en materia de protección de datos.
      </div>
    </div>
    """, unsafe_allow_html=True)
    acepta = st.checkbox("He leído y acepto la cláusula de privacidad *")
    if acepta:
        if st.button("CONTINUAR →", use_container_width=True):
            S.pantalla="datos"; st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# DATOS GENERALES
# ══════════════════════════════════════════════════════════════════════════════
elif S.pantalla == "datos":
    header(True)
    st.markdown('<div class="slabel">Información General del Trabajador</div>',
                unsafe_allow_html=True)
    err = []

    st.markdown('<div class="slabel">Nombre</div>', unsafe_allow_html=True)

    # Nota técnica: Streamlit encapsula sus inputs en shadow DOM, lo que impide
    # que JavaScript externo los controle. La validación se realiza en Python
    # al presionar INICIAR CUESTIONARIO. El CSS text-transform:uppercase
    # da la apariencia visual de mayúsculas; la conversión real ocurre en Python.
    c1,c2,c3 = st.columns(3)
    with c1:
        ap1 = st.text_input("PRIMER APELLIDO", value=S.ap1, key=f"d_ap1_{S.form_v}",
                            help="Solo letras. Sin números ni símbolos.")
    with c2:
        ap2 = st.text_input("SEGUNDO APELLIDO", value=S.ap2, key=f"d_ap2_{S.form_v}",
                            help="Solo letras. Sin números ni símbolos.")
    with c3:
        nom = st.text_input("NOMBRE(S)", value=S.nom, key=f"d_nom_{S.form_v}",
                            help="Solo letras. Sin números ni símbolos.")

    # Validación en tiempo real visible: mostrar error inmediatamente si hay números
    _hay_err_nombre = False
    for _campo, _val, _label in [
        (ap1, ap1, "Primer Apellido"),
        (ap2, ap2, "Segundo Apellido"),
        (nom, nom, "Nombre(s)"),
    ]:
        if _val and re.search(r"[^A-Za-záéíóúÁÉÍÓÚüÜñÑ\s]", _val):
            st.markdown(
                f'<p style="color:#a20000;font-size:.78rem;font-weight:700;margin:0;">' 
                f'⚠ {_label}: solo se permiten letras. Elimina los números o símbolos.</p>',
                unsafe_allow_html=True)
            _hay_err_nombre = True

    c1,c2,c3 = st.columns(3)
    with c1: sexo  = st.selectbox("SEXO",         OPC_SEXO,  key=f"w_sexo_{S.form_v}",   index=idx_de(OPC_SEXO,  S.sexo))
    with c2: edad  = st.selectbox("EDAD (años)",  OPC_EDAD,  key=f"w_edad_{S.form_v}",   index=idx_de(OPC_EDAD,  S.edad))
    with c3: ecivil= st.selectbox("ESTADO CIVIL", OPC_ECIVIL,key=f"w_ecivil_{S.form_v}", index=idx_de(OPC_ECIVIL,S.ecivil))

    st.markdown('<div class="slabel">Nivel de Estudios</div>', unsafe_allow_html=True)
    c1,c2 = st.columns([2,1])
    with c1: estudios = st.selectbox("NIVEL", OPC_ESTUD, key=f"w_estudios_{S.form_v}", index=idx_de(OPC_ESTUD,S.estudios))
    with c2:
        if estudios not in [SEL,"Sin formación"]:
            estatus = st.radio("", ["Terminada","Incompleta"], horizontal=True,
                               key=f"w_estatus_{S.form_v}",
                               index=0 if S.estatus=="Terminada" else 1)
        else:
            estatus="N/A"
            # Si existe la key del radio de estatus la eliminamos para que no persista
            k_est = f"w_estatus_{S.form_v}"
            if k_est in st.session_state: del st.session_state[k_est]

    c1,c2 = st.columns(2)
    with c1: puesto = st.selectbox("PUESTO", OPC_PUESTO, key=f"w_puesto_{S.form_v}", index=idx_de(OPC_PUESTO,S.puesto))
    with c2:
        aopts = [SEL]+S.areas
        area  = st.selectbox("DEPARTAMENTO / ÁREA", aopts, key=f"w_area_{S.form_v}", index=idx_de(aopts,S.area))

    c1,c2 = st.columns(2)
    with c1: contrat  = st.selectbox("TIPO DE CONTRATACIÓN", OPC_CONTRAT,  key=f"w_contrat_{S.form_v}",  index=idx_de(OPC_CONTRAT, S.contrat))
    with c2: personal = st.selectbox("TIPO DE PERSONAL",     OPC_PERSONAL, key=f"w_personal_{S.form_v}", index=idx_de(OPC_PERSONAL,S.personal))

    c1,c2 = st.columns([3,1])
    with c1: jornada = st.selectbox("TIPO DE JORNADA DE TRABAJO", OPC_JORNADA, key=f"w_jornada_{S.form_v}", index=idx_de(OPC_JORNADA,S.jornada))
    with c2:
        st.markdown('<div class="slabel">Rotación de Turnos</div>', unsafe_allow_html=True)
        rotacion = st.radio("", ["Sí","No"], horizontal=True, key=f"w_rotacion_{S.form_v}",
                            index=0 if S.rotacion=="Sí" else 1)

    c1,c2 = st.columns(2)
    with c1: tpuesto = st.selectbox("TIEMPO EN EL PUESTO ACTUAL",   OPC_TPUESTO, key=f"w_tpuesto_{S.form_v}", index=idx_de(OPC_TPUESTO,S.tpuesto))
    with c2: exp     = st.selectbox("TIEMPO DE EXPERIENCIA LABORAL", OPC_EXP,     key=f"w_exp_{S.form_v}",     index=idx_de(OPC_EXP,    S.exp))

    st.markdown("<br>", unsafe_allow_html=True)
    c1,c2 = st.columns([3,1])
    with c1: iniciar = st.button("INICIAR CUESTIONARIO", use_container_width=True)
    with c2:
        if st.button("BORRAR DATOS", use_container_width=True):
            borrar_formulario()
            st.rerun()

    if iniciar:
        # ── Limpiar y validar nombre campo por campo ──────────────────────────
        a1 = solo_letras(ap1.strip())
        a2 = solo_letras(ap2.strip())
        nn = solo_letras(nom.strip())

        # Bloquear si hay caracteres inválidos (ya se muestra el error arriba en tiempo real)
        if _hay_err_nombre:
            err.append("Corrige los campos de nombre antes de continuar: "
                       "solo se permiten letras, sin números ni símbolos.")
        else:
            if not a1: err.append("El Primer Apellido es obligatorio.")
            if not a2: err.append("El Segundo Apellido es obligatorio.")
            if not nn: err.append("El campo Nombre(s) es obligatorio.")

        # ── Resto de campos ───────────────────────────────────────────────────
        if sexo    ==SEL:   err.append("Selecciona el Sexo.")
        if edad    ==SEL:   err.append("Selecciona la Edad.")
        if ecivil  ==SEL:   err.append("Selecciona el Estado Civil.")
        if estudios==SEL:   err.append("Selecciona el Nivel de Estudios.")
        if puesto  ==SEL:   err.append("Selecciona el Puesto.")
        if area    ==SEL:   err.append("Selecciona el Área.")
        if contrat ==SEL:   err.append("Selecciona el Tipo de Contratación.")
        if personal==SEL:   err.append("Selecciona el Tipo de Personal.")
        if jornada ==SEL:   err.append("Selecciona el Tipo de Jornada.")
        if tpuesto ==SEL:   err.append("Selecciona el Tiempo en el Puesto Actual.")
        if exp     ==SEL:   err.append("Selecciona el Tiempo de Experiencia Laboral.")

        if err:
            for e in err:
                st.markdown(
                    f'<p style="color:#a20000;font-size:.82rem;font-weight:600;">' 
                    f'⚠ {e}</p>',
                    unsafe_allow_html=True)
        else:
            # ── Verificar registro duplicado ──────────────────────────────────
            if trabajador_ya_registrado(a1, a2, nn, S.razon, S.cliente_key):
                st.markdown(f"""
                <div class="dup-alert">
                    <div style="font-size:1rem;margin-bottom:.4rem;">
                        ⛔ Registro duplicado detectado
                    </div>
                    El trabajador <strong>{a1} {a2}, {nn}</strong> ya cuenta con un
                    cuestionario registrado en <strong>{S.razon}</strong>.<br><br>
                    Conforme a la NOM-035-STPS-2018, cada trabajador debe ser evaluado
                    una única vez por periodo. Si considera que esto es un error,
                    comuníquese con el responsable del programa.
                </div>
                """, unsafe_allow_html=True)
            else:
                S.ap1=a1; S.ap2=a2; S.nom=nn
                S.sexo=sexo; S.edad=edad; S.ecivil=ecivil
                S.estudios=estudios; S.estatus=estatus
                S.puesto=puesto; S.area=area
                S.contrat=contrat; S.personal=personal
                S.jornada=jornada; S.rotacion=rotacion
                S.tpuesto=tpuesto; S.exp=exp
                S.pantalla="confirmar"; st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# CONFIRMACIÓN
# ══════════════════════════════════════════════════════════════════════════════
elif S.pantalla == "confirmar":
    header(True)
    st.markdown("""
    <div class="cfm-box">
      <p style="font-size:.97rem;font-weight:700;text-transform:uppercase;color:#4b694e;">
        Confirmación de Datos</p>
      <p>Al pulsar el botón <strong>"ACEPTAR"</strong>, usted declara que los datos ingresados
      en el apartado de información general son correctos y reflejan su situación actual
      en la empresa.</p>
    </div>
    """, unsafe_allow_html=True)
    c1,c2 = st.columns(2)
    with c1:
        if st.button("ACEPTAR", use_container_width=True):
            S.pantalla="aviso"; st.rerun()
    with c2:
        if st.button("← REGRESAR Y EDITAR", use_container_width=True):
            S.pantalla="datos"; st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# AVISO DE NO RETROCESO
# ══════════════════════════════════════════════════════════════════════════════
elif S.pantalla == "aviso":
    header(True)
    st.markdown("""
    <div class="av-box">
      <div class="av-tit">⚠ ANTES DE COMENZAR — LEA CON ATENCIÓN</div>
      <p>Este cuestionario está diseñado para capturar su percepción inmediata.
      Una vez que avance a la siguiente pregunta,
      <strong>no podrá regresar a la anterior.</strong><br>
      Por favor, responda con total sinceridad basándose en su primera impresión.</p>
    </div>
    """, unsafe_allow_html=True)
    if st.button("COMENZAR CUESTIONARIO →", use_container_width=True):
        reset_cuestionario()
        S.pantalla="preguntas"; st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# PREGUNTAS — una a la vez, sin retroceso
# ══════════════════════════════════════════════════════════════════════════════
elif S.pantalla == "preguntas":
    header(True)

    # Garantizar listas reales (nunca None)
    if not isinstance(S.r1, list): S.r1 = []
    if not isinstance(S.r2, list): S.r2 = []
    if not isinstance(S.r3, list): S.r3 = []
    if not isinstance(S.r4, list): S.r4 = []

    sec_obj = SECCIONES[S.sec]
    preg    = sec_obj["preg"]
    idx     = S.preg

    st.markdown(
        f'<div class="prog-txt">{sec_obj["titulo"]} &nbsp;·&nbsp; '
        f'Pregunta {idx + 1} de {len(preg)}</div>',
        unsafe_allow_html=True)
    st.progress(idx / len(preg))

    st.markdown(f"""
    <div class="pq-card">
      <div class="pq-sec">{sec_obj["titulo"]}</div>
      <div class="pq-ins">{sec_obj["inst"]}</div>
      <div class="pq-txt">{preg[idx]}</div>
    </div>
    """, unsafe_allow_html=True)

    resp = st.radio(
        "", ["Sí", "No"],
        index=None, horizontal=True,
        key=f"q_{S.sec}_{idx}",
        label_visibility="collapsed",
    )

    if S.err:
        st.markdown('<p class="err-r">⚠ SELECCIONA UNA RESPUESTA ANTES DE CONTINUAR</p>',
                    unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    if st.button("SIGUIENTE →", use_container_width=True, key=f"sig_{S.sec}_{idx}"):

        if resp is None:
            S.err = True
            st.rerun()

        else:
            S.err = False

            # ── Guardar con concatenación (más confiable que append en session_state)
            if   S.sec == 0: S.r1 = S.r1 + [resp]
            elif S.sec == 1: S.r2 = S.r2 + [resp]
            elif S.sec == 2: S.r3 = S.r3 + [resp]
            elif S.sec == 3: S.r4 = S.r4 + [resp]

            if (idx + 1) < len(preg):
                # Siguiente pregunta de la misma sección
                S.preg = idx + 1
                st.rerun()
            else:
                # ── Fin de sección ────────────────────────────────────────────
                if S.sec == 0:
                    if all(r == "No" for r in S.r1):
                        resultado = evaluar(S.r1, [], [], [])
                        S.res = resultado
                        guardar(dict(
                            folio=S.folio, cliente=S.cliente_key, razon=S.razon,
                            nombre=f"{S.ap1}; {S.ap2}; {S.nom}",
                            sexo=S.sexo, edad=S.edad, ecivil=S.ecivil,
                            estudios=S.estudios, estatus=S.estatus,
                            puesto=S.puesto, area=S.area,
                            contrat=S.contrat, personal=S.personal,
                            jornada=S.jornada, rotacion=S.rotacion,
                            tpuesto=S.tpuesto, exp=S.exp,
                            r1=S.r1, r2=[], r3=[], r4=[], res=resultado))
                        S.pantalla = "fin"
                    else:
                        S.sec = 1; S.preg = 0

                elif S.sec == 1:
                    S.sec = 2; S.preg = 0

                elif S.sec == 2:
                    S.sec = 3; S.preg = 0

                elif S.sec == 3:
                    resultado = evaluar(S.r1, S.r2, S.r3, S.r4)
                    S.res = resultado
                    guardar(dict(
                        folio=S.folio, cliente=S.cliente_key, razon=S.razon,
                        nombre=f"{S.ap1}; {S.ap2}; {S.nom}",
                        sexo=S.sexo, edad=S.edad, ecivil=S.ecivil,
                        estudios=S.estudios, estatus=S.estatus,
                        puesto=S.puesto, area=S.area,
                        contrat=S.contrat, personal=S.personal,
                        jornada=S.jornada, rotacion=S.rotacion,
                        tpuesto=S.tpuesto, exp=S.exp,
                        r1=S.r1, r2=S.r2, r3=S.r3, r4=S.r4, res=resultado))
                    S.pantalla = "fin"

                st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# FIN
# ══════════════════════════════════════════════════════════════════════════════
elif S.pantalla == "fin":
    header(True)
    st.markdown("""
    <div class="fin-box">
      <div class="fin-tit">✓ CUESTIONARIO FINALIZADO CORRECTAMENTE</div>
      <div class="fin-sub">AGRADECEMOS TU PARTICIPACIÓN</div>
    </div>
    """, unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)

    # Solo el operativo puede ver los resultados
    if not _MODO_EMPLEADO:
        with st.expander("Ver resultado — uso interno operativo"):
            if S.res:
                r = S.res
                cls = "res-r" if r["atencion"] else "res-v"
                st.markdown(f"""
                <div class="{cls}">
                  <div style="font-size:.7rem;opacity:.75;text-transform:uppercase;letter-spacing:.1em;">
                    Resultado · Guía I · NOM-035-STPS-2018</div>
                  <div class="res-niv">{r['nivel']}</div>
                  <div class="res-sub">S1: {r['p1']} &nbsp;|&nbsp; S2: {r['p2']} &nbsp;|&nbsp;
                    S3: {r['p3']} &nbsp;|&nbsp; S4: {r['p4']}
                  </div>
                </div>
                <div class="ac-box"><strong>Descripción:</strong><br>{r['desc']}</div>
                """, unsafe_allow_html=True)
                if r["criterios"] != "Ninguno":
                    st.markdown(f'<div class="dt-box"><strong>Criterios activados:</strong><br>{r["criterios"]}</div>',
                                unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    if _MODO_EMPLEADO:
        st.markdown("""
        <div style="text-align:center;font-size:.92rem;color:#666;
                    font-family:Montserrat,sans-serif;padding:1rem 0;">
            Puedes cerrar esta ventana.
        </div>
        """, unsafe_allow_html=True)
    else:
        c1,c2 = st.columns(2)
        with c1:
            if st.button("REGISTRAR OTRO EMPLEADO", use_container_width=True):
                S.pantalla="panel"; st.rerun()
        with c2:
            if st.button("VOLVER AL PANEL", use_container_width=True):
                S.pantalla="panel"; st.rerun()
